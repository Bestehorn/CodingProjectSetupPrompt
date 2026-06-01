"""Regenerable generator for the CV Customizer test fixtures (.docx only).

This script produces the binary ``.docx`` fixtures used by the suite's
deterministic Python tests. Plain-text fixtures (``sample_jd.html``,
``sample_jd.txt``, ``sample_database.md``) are authored directly and are not
generated here.

Run from anywhere:

    python cli-agents/cv/tests/fixtures/make_fixtures.py

Outputs are written next to this script (the versioned fixtures directory),
NOT into the gitignored ``tmp/`` tree, so they are committed and reproducible.

Dependency policy (matches design): this script does NOT install anything. If
``python-docx`` is unavailable it exits non-zero with a clear message naming
the missing package, so a human can install it once. Agents must not run pip.
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ModuleNotFoundError as exc:  # pragma: no cover - exercised only when dep missing
    missing = getattr(exc, "name", "python-docx")
    sys.stderr.write(
        f"ERROR: required package for fixture generation is not installed: {missing}.\n"
        "Install 'python-docx' in this Python environment and re-run. "
        "This script does not install packages.\n"
    )
    raise SystemExit(2)

FIXTURES_DIR = Path(__file__).resolve().parent

# A block of filler prose sized to comfortably fill roughly one rendered page
# alongside a heading. Calibration is approximate here; the authoritative
# 1/2/3-page check runs on a Word-equipped host in task 18.
_PARAGRAPH = (
    "This paragraph exists to occupy vertical space so the rendered document "
    "reaches a predictable length. It repeats ordinary sentences that carry no "
    "special meaning. The fixtures rely on explicit hard page breaks to make "
    "the page count deterministic regardless of the rendering engine, while "
    "this filler text ensures each page also carries visible body content. "
)


def _fill_page(doc: Document, page_number: int, total_pages: int) -> None:
    """Add a heading plus filler paragraphs representing one page of content."""
    doc.add_heading(f"Calibration Page {page_number} of {total_pages}", level=1)
    for _ in range(3):
        doc.add_paragraph(_PARAGRAPH)


def make_calibrated(path: Path, pages: int) -> None:
    """Create a doc that renders to exactly ``pages`` pages via hard breaks."""
    doc = Document()
    for page_number in range(1, pages + 1):
        if page_number > 1:
            doc.add_page_break()
        _fill_page(doc, page_number, pages)
    doc.save(str(path))


def _set_columns(section, num: int) -> None:
    """Force a section to use ``num`` newspaper-style columns (ATS hazard)."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))


def _add_text_box(doc: Document, text: str) -> None:
    """Append a paragraph containing a VML text box (``w:txbxContent``)."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    v_ns = "urn:schemas-microsoft-com:vml"
    run_xml = (
        '<w:r xmlns:w="%s" xmlns:v="%s">'
        "<w:pict>"
        '<v:shape type="#_x0000_t202" style="width:240pt;height:48pt">'
        "<v:textbox>"
        "<w:txbxContent>"
        "<w:p><w:r><w:t>%s</w:t></w:r></w:p>"
        "</w:txbxContent>"
        "</v:textbox>"
        "</v:shape>"
        "</w:pict>"
        "</w:r>"
    ) % (w_ns, v_ns, text)
    paragraph = doc.add_paragraph()
    paragraph._p.append(parse_xml(run_xml))


def make_ats_hazards(path: Path) -> None:
    """Create a doc bundling several ATS-hostile structures.

    Contains: header content, a text box (``w:txbxContent``), a 2-column
    section, and a layout table.
    """
    doc = Document()

    # Header content (ATS parsers commonly ignore headers).
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Jordan Candidate  |  +1 555 0100  |  jordan@example.com"

    doc.add_heading("ATS Hazard Sample", level=1)
    doc.add_paragraph(
        "This document intentionally contains structures that automated "
        "applicant tracking systems frequently mis-parse."
    )

    # Text box hazard.
    _add_text_box(doc, "Key skills are hidden inside this text box (ATS hazard).")

    # Layout table hazard.
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Experience"
    table.cell(0, 1).text = "Skills"
    table.cell(1, 0).text = "Built data pipelines and services."
    table.cell(1, 1).text = "Python, SQL, AWS"

    # Two-column section hazard.
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_columns(section, 2)
    doc.add_paragraph(
        "This passage sits in a two-column layout that ATS parsers often read "
        "out of order, scrambling the reading sequence of the content."
    )
    doc.add_paragraph(
        "A second two-column paragraph reinforces the multi-column hazard so "
        "the structural checker has clear content to detect."
    )

    doc.save(str(path))


def make_sample_cv(path: Path) -> None:
    """Create a small, clean, single-column sample CV."""
    doc = Document()

    title = doc.add_heading("Alex Morgan", level=0)
    title.runs[0].font.size = Pt(20)
    doc.add_paragraph("Software Engineer  |  alex.morgan@example.com  |  +1 555 0123")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Software engineer with five years of experience building backend "
        "services and data pipelines. Comfortable across the stack and focused "
        "on reliable, well-tested systems."
    )

    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph("Senior Software Engineer, Northwind Labs (2021-present)")
    for bullet in (
        "Designed and shipped a service that processes one million events per day.",
        "Reduced API latency by 40 percent through query and caching improvements.",
        "Mentored two junior engineers and led the team's code-review practice.",
    ):
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("Software Engineer, Acme Corp (2019-2021)")
    for bullet in (
        "Built internal tooling in Python that automated weekly reporting.",
        "Migrated a monolith module to a documented REST API.",
    ):
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Sc. in Computer Science, State University (2019)")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, SQL, AWS, Docker, REST APIs, testing")

    doc.save(str(path))


def main() -> int:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    make_calibrated(FIXTURES_DIR / "page_1.docx", pages=1)
    make_calibrated(FIXTURES_DIR / "page_2.docx", pages=2)
    make_calibrated(FIXTURES_DIR / "page_3.docx", pages=3)
    make_ats_hazards(FIXTURES_DIR / "ats_hazards.docx")
    make_sample_cv(FIXTURES_DIR / "sample_cv.docx")
    generated = [
        "page_1.docx",
        "page_2.docx",
        "page_3.docx",
        "ats_hazards.docx",
        "sample_cv.docx",
    ]
    print("Generated .docx fixtures in", FIXTURES_DIR)
    for name in generated:
        print("  -", name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
