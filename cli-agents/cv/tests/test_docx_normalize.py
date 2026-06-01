"""Unit tests for ``docx_normalize.py`` (task 2 / subtask 2.1).

Covers the stable-anchor contract from the design's "anchor model":

* keys are stable after inserting/deleting *earlier* paragraphs (a key computed
  on the original doc still resolves to the same logical paragraph);
* duplicate (identical) paragraphs get distinct ordinals;
* every key resolves to exactly one paragraph;
* anchors round-trip against the live doc (load ``*.anchors.json``, resolve each
  key, confirm it maps back to the right paragraph).

Tests build small in-memory docs with python-docx and also exercise the
versioned ``sample_cv.docx`` fixture from task 1. No environment variables are
used; ``docx_normalize`` is importable via the path wiring in ``conftest.py``.
"""

from __future__ import annotations

import json
import subprocess
import sys
import textwrap
from pathlib import Path

import pytest
from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

import docx_normalize as dn


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------


def build_doc(items):
    """Build an in-memory Document from a list of ``(kind, text)`` tuples.

    kind: ``"h"`` heading (Heading 1), ``"p"`` body paragraph, ``"b"`` bullet.
    A fresh ``Document()`` starts empty, so the items become the body in order.
    """
    doc = Document()
    for kind, text in items:
        if kind == "h":
            doc.add_heading(text, level=1)
        elif kind == "b":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    return doc


def delete_paragraph(paragraph) -> None:
    """Remove a paragraph from its parent body (python-docx has no public API)."""
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def key_of(doc, text):
    """Return the stable key of the first paragraph whose text == ``text``."""
    for anchor in dn.compute_paragraph_anchors(doc):
        if anchor.text == text:
            return anchor.key
    raise AssertionError(f"no paragraph with text {text!r}")


SAMPLE_ITEMS = [
    ("h", "Professional Summary"),
    ("p", "Engineer with experience building services."),
    ("h", "Professional Experience"),
    ("p", "Senior Engineer, Northwind Labs (2021-present)"),
    ("b", "Designed a service handling a million events per day."),
    ("b", "Reduced latency by 40 percent."),
    ("h", "Education"),
    ("p", "B.Sc. Computer Science, State University (2019)"),
]


# --------------------------------------------------------------------------
# duplicate ordinals
# --------------------------------------------------------------------------


def test_duplicate_paragraphs_get_distinct_ordinals():
    doc = build_doc(
        [
            ("h", "Experience"),
            ("p", "Delivered measurable results."),
            ("p", "Delivered measurable results."),
            ("p", "Delivered measurable results."),
        ]
    )
    anchors = [a for a in dn.compute_paragraph_anchors(doc) if not a.is_heading]
    keys = [a.key for a in anchors]
    # Three identical paragraphs -> three distinct keys (ordinals 0, 1, 2).
    assert len(set(keys)) == 3
    ordinals = sorted(int(k.rsplit("::", 1)[1]) for k in keys)
    assert ordinals == [0, 1, 2]
    # Each duplicate still resolves to exactly one (its own) paragraph.
    for anchor in anchors:
        assert dn.resolve_paragraph_index(doc, anchor.key) == anchor.index


def test_identical_paragraphs_in_different_sections_have_distinct_keys():
    doc = build_doc(
        [
            ("h", "Summary"),
            ("p", "Results-oriented professional."),
            ("h", "Profile"),
            ("p", "Results-oriented professional."),
        ]
    )
    bodies = [a for a in dn.compute_paragraph_anchors(doc) if not a.is_heading]
    assert bodies[0].key != bodies[1].key  # different section token
    assert bodies[0].section == "Summary"
    assert bodies[1].section == "Profile"


# --------------------------------------------------------------------------
# every key resolves to exactly one paragraph
# --------------------------------------------------------------------------


def test_every_key_resolves_to_exactly_one_paragraph_in_memory():
    doc = build_doc(SAMPLE_ITEMS)
    anchors = dn.compute_paragraph_anchors(doc)
    mapping = dn.compute_anchor_map(doc)
    # No key collisions: one entry per paragraph.
    assert len(mapping) == len(doc.paragraphs) == len(anchors)
    for anchor in anchors:
        assert dn.resolve_paragraph_index(doc, anchor.key) == anchor.index
        # resolve_paragraph_key returns the same underlying paragraph object.
        assert dn.resolve_paragraph_key(doc, anchor.key).text == anchor.text


def test_every_key_resolves_to_exactly_one_paragraph_sample_cv(fixture_path):
    doc = Document(str(fixture_path("sample_cv.docx")))
    anchors = dn.compute_paragraph_anchors(doc)
    mapping = dn.compute_anchor_map(doc)
    assert len(mapping) == len(doc.paragraphs)
    for anchor in anchors:
        assert dn.resolve_paragraph_index(doc, anchor.key) == anchor.index


def test_unresolvable_key_returns_none_and_raises():
    doc = build_doc(SAMPLE_ITEMS)
    missing = dn.make_paragraph_key("Nonexistent Section", "never appears here", 0)
    assert dn.resolve_paragraph_index(doc, missing) is None
    with pytest.raises(dn.AnchorResolutionError):
        dn.resolve_paragraph_key(doc, missing)


def test_editing_a_paragraphs_own_text_invalidates_its_key():
    doc = build_doc(SAMPLE_ITEMS)
    target_text = "B.Sc. Computer Science, State University (2019)"
    original_key = key_of(doc, target_text)
    # Mutate the paragraph's own content -> its content hash changes.
    para = dn.resolve_paragraph_key(doc, original_key)
    para.text = "M.Sc. Computer Science, State University (2021)"
    # The old key no longer resolves (caller must report unresolved, never guess).
    assert dn.resolve_paragraph_index(doc, original_key) is None


# --------------------------------------------------------------------------
# stability under inserting / deleting EARLIER paragraphs
# --------------------------------------------------------------------------


def test_key_stable_after_inserting_earlier_paragraph():
    doc = build_doc(SAMPLE_ITEMS)
    target_text = "B.Sc. Computer Science, State University (2019)"
    original_key = key_of(doc, target_text)
    before_count = len(doc.paragraphs)

    # Insert a brand-new, distinct paragraph in an EARLIER section.
    first_body = doc.paragraphs[1]  # "Engineer with experience..."
    first_body.insert_paragraph_before("An additional earlier line of context.")

    assert len(doc.paragraphs) == before_count + 1
    # The original key still resolves to exactly one paragraph: the same one.
    resolved_index = dn.resolve_paragraph_index(doc, original_key)
    assert resolved_index is not None
    assert doc.paragraphs[resolved_index].text == target_text
    # And the key recomputed on the edited doc is unchanged.
    assert key_of(doc, target_text) == original_key


def test_key_stable_after_deleting_earlier_paragraph():
    doc = build_doc(SAMPLE_ITEMS)
    target_text = "B.Sc. Computer Science, State University (2019)"
    original_key = key_of(doc, target_text)
    before_count = len(doc.paragraphs)

    # Delete a distinct EARLIER paragraph (a bullet in Professional Experience).
    victim = next(
        p for p in doc.paragraphs if p.text == "Reduced latency by 40 percent."
    )
    delete_paragraph(victim)

    assert len(doc.paragraphs) == before_count - 1
    resolved_index = dn.resolve_paragraph_index(doc, original_key)
    assert resolved_index is not None
    assert doc.paragraphs[resolved_index].text == target_text
    assert key_of(doc, target_text) == original_key


def test_key_stable_after_inserting_and_deleting_combined():
    doc = build_doc(SAMPLE_ITEMS)
    target_text = "Senior Engineer, Northwind Labs (2021-present)"
    original_key = key_of(doc, target_text)

    # Insert earlier, then delete a different earlier paragraph.
    doc.paragraphs[1].insert_paragraph_before("Yet another earlier paragraph.")
    victim = next(
        p for p in doc.paragraphs if p.text == "Engineer with experience building services."
    )
    delete_paragraph(victim)

    resolved_index = dn.resolve_paragraph_index(doc, original_key)
    assert resolved_index is not None
    assert doc.paragraphs[resolved_index].text == target_text


# --------------------------------------------------------------------------
# anchors.json round-trip against the live doc
# --------------------------------------------------------------------------


def test_anchors_json_round_trip(fixture_path, tmp_path):
    out_md, out_anchors = dn.normalize_docx(
        fixture_path("sample_cv.docx"), tmp_path / "cv.normalized.md"
    )
    # Default sidecar naming matches the design convention.
    assert out_anchors == tmp_path / "cv.anchors.json"
    assert out_md.exists() and out_anchors.exists()

    sidecar = dn.load_anchors(out_anchors)
    assert sidecar["schema"] == dn.ANCHORS_SCHEMA
    mapping = sidecar["anchors"]
    paragraphs_meta = {p["index"]: p for p in sidecar["paragraphs"]}

    # Re-open the document as a separate live object and resolve every key.
    live = Document(str(fixture_path("sample_cv.docx")))
    assert len(mapping) == len(live.paragraphs)
    seen_indices = set()
    for key, index in mapping.items():
        resolved = dn.resolve_paragraph_index(live, key)
        assert resolved == index, f"key {key!r} resolved to {resolved}, expected {index}"
        assert resolved not in seen_indices, "two keys resolved to the same paragraph"
        seen_indices.add(resolved)
        # The stored text matches the live paragraph at that index.
        assert live.paragraphs[index].text == paragraphs_meta[index]["text"]
    # Every paragraph was covered exactly once.
    assert seen_indices == set(range(len(live.paragraphs)))


def test_anchors_round_trip_holds_after_earlier_edit(fixture_path, tmp_path):
    """A key from the original sidecar still resolves after an earlier edit."""
    _, out_anchors = dn.normalize_docx(
        fixture_path("sample_cv.docx"), tmp_path / "cv.normalized.md"
    )
    sidecar = dn.load_anchors(out_anchors)

    live = Document(str(fixture_path("sample_cv.docx")))
    # Pick a late paragraph to track (the Skills line near the end).
    skills_text = "Python, SQL, AWS, Docker, REST APIs, testing"
    tracked_key = key_of(live, skills_text)
    assert tracked_key in sidecar["anchors"]

    # Apply an earlier edit: insert a new bullet under Professional Experience.
    exp_para = next(
        p for p in live.paragraphs if p.text == "Senior Software Engineer, Northwind Labs (2021-present)"
    )
    exp_para.insert_paragraph_before("Led an additional cross-team initiative.", style="List Bullet")

    resolved = dn.resolve_paragraph_index(live, tracked_key)
    assert resolved is not None
    assert live.paragraphs[resolved].text == skills_text


# --------------------------------------------------------------------------
# sidecar path derivation + markdown surface
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "output_md, expected",
    [
        ("cv.normalized.md", "cv.anchors.json"),
        ("letter.normalized.md", "letter.anchors.json"),
        ("foo.md", "foo.anchors.json"),
        ("plain", "plain.anchors.json"),
    ],
)
def test_anchors_path_for_naming(output_md, expected, tmp_path):
    assert dn.anchors_path_for(tmp_path / output_md) == tmp_path / expected


def test_to_markdown_includes_headings_and_bullets(fixture_path):
    md = dn.to_markdown(Document(str(fixture_path("sample_cv.docx"))))
    assert "# Alex Morgan" in md  # Title -> level 0 -> single '#'
    assert "## Professional Experience" in md  # Heading 1 -> '##'
    assert "- Designed and shipped a service" in md  # List Bullet -> '- '


# --------------------------------------------------------------------------
# CLI behaviour
# --------------------------------------------------------------------------


def test_cli_writes_md_and_anchors(fixture_path, tmp_path, scripts_dir):
    out_md = tmp_path / "cv.normalized.md"
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "docx_normalize.py"),
            str(fixture_path("sample_cv.docx")),
            str(out_md),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert Path(payload["normalized_md"]).exists()
    assert Path(payload["anchors_json"]) == tmp_path / "cv.anchors.json"
    assert Path(payload["anchors_json"]).exists()


def test_cli_missing_input_exits_nonzero(tmp_path, scripts_dir):
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "docx_normalize.py"),
            str(tmp_path / "does_not_exist.docx"),
            str(tmp_path / "out.md"),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 2
    assert "not found" in result.stderr.lower()


def test_missing_python_docx_exits_nonzero_naming_package(tmp_path, scripts_dir):
    """If python-docx cannot be imported, the script exits 2 naming the package."""
    blocker = tmp_path / "blocker.py"
    blocker.write_text(
        textwrap.dedent(
            f"""
            import sys, importlib.abc

            class _Block(importlib.abc.MetaPathFinder):
                def find_spec(self, name, path, target=None):
                    if name == "docx" or name.startswith("docx."):
                        raise ModuleNotFoundError(f"No module named {{name!r}}", name=name)
                    return None

            sys.meta_path.insert(0, _Block())
            sys.path.insert(0, {str(scripts_dir)!r})
            import docx_normalize  # noqa: F401  (import triggers the guarded failure)
            """
        ),
        encoding="utf-8",
    )
    result = subprocess.run(
        [sys.executable, str(blocker)], capture_output=True, text=True
    )
    assert result.returncode == 2
    assert "python-docx" in result.stderr


# --------------------------------------------------------------------------
# property-based coverage (random in-memory docs)
# --------------------------------------------------------------------------

_TEXT_POOL = st.sampled_from(
    ["Alpha", "Beta", "Gamma", "Delta", "Repeated line", "Repeated line"]
)
_ITEM = st.tuples(st.sampled_from(["h", "p", "b"]), _TEXT_POOL)


@settings(max_examples=60, deadline=None)
@given(items=st.lists(_ITEM, min_size=1, max_size=14))
def test_property_keys_unique_and_resolve(items):
    doc = build_doc(items)
    anchors = dn.compute_paragraph_anchors(doc)
    keys = [a.key for a in anchors]
    assert len(set(keys)) == len(keys), "keys must be unique (distinct ordinals)"
    mapping = dn.compute_anchor_map(doc)
    assert len(mapping) == len(doc.paragraphs)
    for anchor in anchors:
        assert dn.resolve_paragraph_index(doc, anchor.key) == anchor.index


@settings(max_examples=60, deadline=None)
@given(
    items=st.lists(_ITEM, min_size=1, max_size=14),
    data=st.data(),
)
def test_property_keys_stable_under_unique_earlier_insert(items, data):
    doc = build_doc(items)
    before = {a.key: a.text for a in dn.compute_paragraph_anchors(doc)}

    # Insert a globally-unique paragraph before a randomly chosen paragraph.
    idx = data.draw(st.integers(min_value=0, max_value=len(doc.paragraphs) - 1))
    doc.paragraphs[idx].insert_paragraph_before("GLOBALLY-UNIQUE-INSERT-7f3a9c2e")

    # Every previously computed key still resolves to exactly its paragraph.
    for key, text in before.items():
        resolved = dn.resolve_paragraph_index(doc, key)
        assert resolved is not None, f"key {key!r} stopped resolving after insert"
        assert doc.paragraphs[resolved].text == text
