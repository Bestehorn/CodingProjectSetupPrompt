"""Sanity tests for the versioned test fixtures.

These confirm the fixture set is present and well-formed so later tasks (the
deterministic Python core in tasks 2-6) can rely on them. The authoritative
1/2/3-page rendered page count is validated separately on a Word-equipped host
(task 18); here we only assert structural properties that hold without a
renderer.
"""

from __future__ import annotations

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx(fixture_path, name):
    return Document(str(fixture_path(name)))


def test_text_fixtures_present_and_nonempty(fixture_path):
    for name in ("sample_jd.html", "sample_jd.txt", "sample_database.md"):
        text = fixture_path(name).read_text(encoding="utf-8")
        assert text.strip(), f"{name} should not be empty"


def test_jd_fixtures_share_core_requirements(fixture_path):
    html = fixture_path("sample_jd.html").read_text(encoding="utf-8")
    txt = fixture_path("sample_jd.txt").read_text(encoding="utf-8")
    for token in ("Python", "SQL", "AWS"):
        assert token in html
        assert token in txt


def test_database_markdown_has_sections(fixture_path):
    md = fixture_path("sample_database.md").read_text(encoding="utf-8")
    assert "## Skills" in md
    assert "## Achievements" in md


def test_calibrated_docx_have_expected_hard_page_breaks(fixture_path):
    # page_N.docx uses N-1 explicit hard page breaks plus body content per page.
    for name, pages in (("page_1.docx", 1), ("page_2.docx", 2), ("page_3.docx", 3)):
        doc = _docx(fixture_path, name)
        break_count = len(doc.element.findall(f".//{{{W_NS}}}br[@{{{W_NS}}}type='page']")) + len(
            doc.element.findall(f".//{{{W_NS}}}lastRenderedPageBreak")
        )
        # python-docx renders add_page_break() as a run-level <w:br w:type="page"/>.
        hard_breaks = len(
            doc.element.findall(f".//{{{W_NS}}}br[@{{{W_NS}}}type='page']")
        )
        assert hard_breaks == pages - 1, (
            f"{name} expected {pages - 1} hard page break(s), found {hard_breaks}"
        )
        assert break_count >= pages - 1


def test_sample_cv_is_clean_single_column(fixture_path):
    doc = _docx(fixture_path, "sample_cv.docx")
    # No multi-column sections.
    for cols in doc.element.iter(f"{{{W_NS}}}cols"):
        num = cols.get(f"{{{W_NS}}}num")
        assert num is None or int(num) <= 1
    # No tables, no text boxes (it is the "clean" doc).
    assert len(doc.tables) == 0
    assert "txbxContent" not in doc.element.xml
    # Has the expected headings.
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Professional Experience" in headings


def test_ats_hazards_docx_contains_all_hazards(fixture_path):
    doc = _docx(fixture_path, "ats_hazards.docx")
    # Text box.
    assert "txbxContent" in doc.element.xml
    # Multi-column section.
    has_multicol = any(
        (cols.get(f"{{{W_NS}}}num") is not None and int(cols.get(f"{{{W_NS}}}num")) > 1)
        for cols in doc.element.iter(f"{{{W_NS}}}cols")
    )
    assert has_multicol, "expected a section with more than one column"
    # Layout table.
    assert len(doc.tables) >= 1
    # Header content.
    assert doc.sections[0].header.paragraphs[0].text.strip()
