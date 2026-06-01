"""Unit tests for ``docx_edit.py`` (task 5 / subtask 5.1).

Covers the edit engine from the design's "Edit-Application Layer":

* each operation in the closed vocabulary (``replace_run_text``,
  ``replace_paragraph_text``, ``insert_paragraph_after``/``before``,
  ``delete_paragraph``, ``set_paragraph_style``, ``replace_bullet_list``)
  applied to in-memory docs and the versioned ``sample_cv.docx`` fixture;
* run-formatting preserved on an intra-run replacement, and
  ``formatting_normalized`` recorded when a match spans runs;
* idempotent re-run (a second application of the same Change_List marks entries
  ``already_satisfied`` and does not corrupt the document);
* ``failed_to_apply`` when an anchor cannot be resolved, asserting no OTHER
  paragraph changed (Property 9);
* verification output correctness (the ``result.json`` statuses match what
  actually happened on disk).

Tests build small in-memory docs with python-docx, saving them to ``tmp_path``
because the engine re-reads the document from disk during verification. No
environment variables are used; ``docx_edit``/``docx_normalize`` are importable
via the path wiring in ``conftest.py``.
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
import textwrap
from pathlib import Path

import pytest
from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

import docx_edit as de
import docx_normalize as dn


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------


def key_for_text(doc, text):
    """Stable paragraph_key of the first paragraph whose text == ``text``."""
    for anchor in dn.compute_paragraph_anchors(doc):
        if anchor.text == text:
            return anchor.key
    raise AssertionError(f"no paragraph with text {text!r}")


def save(doc, path: Path) -> Path:
    doc.save(str(path))
    return path


def texts(path: Path):
    """All body paragraph texts of the document at ``path`` (order preserved)."""
    return [p.text for p in Document(str(path)).paragraphs]


def status_by_id(result: dict) -> dict:
    return {e["id"]: e["status"] for e in result["entries"]}


def build_simple(path: Path) -> Path:
    """A small clean doc: one heading + three distinct body paragraphs."""
    doc = Document()
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Alpha line of content.")
    doc.add_paragraph("Beta line of content.")
    doc.add_paragraph("Gamma line of content.")
    return save(doc, path)


def build_bulleted(path: Path) -> Path:
    """A heading, three bullets, then a trailing non-list paragraph."""
    doc = Document()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("First bullet item.", style="List Bullet")
    doc.add_paragraph("Second bullet item.", style="List Bullet")
    doc.add_paragraph("Third bullet item.", style="List Bullet")
    doc.add_paragraph("Trailing non-list paragraph.")
    return save(doc, path)


# --------------------------------------------------------------------------
# replace_run_text
# --------------------------------------------------------------------------


def test_replace_run_text_basic(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Beta line of content.")
    entry = {
        "id": "CL-1",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "Beta"},
        "new_text": "Delta",
        "implements_findings": ["F-1"],
    }
    result = de.apply_change_list(path, [entry], iteration=1, write_result=False)
    assert status_by_id(result)["CL-1"] == "verified"
    assert "Delta line of content." in texts(path)
    assert "Beta line of content." not in texts(path)


def test_replace_run_text_preserves_intra_run_formatting(tmp_path):
    """A bold run stays bold after an intra-run replacement (no normalization)."""
    doc = Document()
    doc.add_heading("Header", level=1)
    para = doc.add_paragraph()
    run = para.add_run("Hello World")
    run.bold = True
    run.italic = True
    path = save(doc, tmp_path / "fmt.docx")

    key = key_for_text(doc, "Hello World")
    entry = {
        "id": "CL-fmt",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "World"},
        "new_text": "Universe",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-fmt"] == "verified"  # not formatting_normalized

    edited = Document(str(path))
    target = next(p for p in edited.paragraphs if p.text == "Hello Universe")
    assert len(target.runs) == 1
    assert target.runs[0].bold is True
    assert target.runs[0].italic is True


def test_replace_run_text_cross_run_records_formatting_normalized(tmp_path):
    """A match spanning runs flattens to one run and is flagged normalized."""
    doc = Document()
    doc.add_heading("Header", level=1)
    para = doc.add_paragraph()
    r0 = para.add_run("Hello ")
    r0.bold = True
    r1 = para.add_run("brave ")  # plain
    r2 = para.add_run("World")
    r2.italic = True
    path = save(doc, tmp_path / "span.docx")

    key = key_for_text(doc, "Hello brave World")
    entry = {
        "id": "CL-span",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "brave World"},
        "new_text": "everyone",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    record = result["entries"][0]
    assert record["status"] == "formatting_normalized"
    assert record["formatting_normalized"] is True
    assert record["applied"] is True

    edited = Document(str(path))
    target = next(p for p in edited.paragraphs if p.text == "Hello everyone")
    # The two trailing runs collapsed into the first matched run's formatting.
    assert len(target.runs) == 2  # "Hello " (bold) + "everyone" (from first matched run)


def test_replace_run_text_match_absent_is_failed(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {
        "id": "CL-x",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "Nonexistent"},
        "new_text": "whatever",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-x"] == "failed_to_apply"
    # The paragraph is untouched.
    assert "Alpha line of content." in texts(path)


# --------------------------------------------------------------------------
# replace_paragraph_text
# --------------------------------------------------------------------------


def test_replace_paragraph_text(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Gamma line of content.")
    entry = {
        "id": "CL-p",
        "operation": "replace_paragraph_text",
        "anchor": {"paragraph_key": key},
        "new_text": "A fully rewritten paragraph.",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-p"] == "verified"
    assert "A fully rewritten paragraph." in texts(path)
    assert "Gamma line of content." not in texts(path)


# --------------------------------------------------------------------------
# insert_paragraph_after / insert_paragraph_before
# --------------------------------------------------------------------------


def test_insert_paragraph_after(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {
        "id": "CL-ia",
        "operation": "insert_paragraph_after",
        "anchor": {"paragraph_key": key},
        "new_text": "Inserted right after Alpha.",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-ia"] == "verified"
    body = texts(path)
    assert body.index("Inserted right after Alpha.") == body.index("Alpha line of content.") + 1


def test_insert_paragraph_before(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Beta line of content.")
    entry = {
        "id": "CL-ib",
        "operation": "insert_paragraph_before",
        "anchor": {"paragraph_key": key},
        "new_text": "Inserted right before Beta.",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-ib"] == "verified"
    body = texts(path)
    assert body.index("Inserted right before Beta.") == body.index("Beta line of content.") - 1


def test_insert_paragraph_after_copies_anchor_style_by_default(tmp_path):
    path = build_bulleted(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "First bullet item.")
    entry = {
        "id": "CL-istyle",
        "operation": "insert_paragraph_after",
        "anchor": {"paragraph_key": key},
        "new_text": "A fourth bullet inserted after the first.",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-istyle"] == "verified"
    edited = Document(str(path))
    inserted = next(p for p in edited.paragraphs if p.text == "A fourth bullet inserted after the first.")
    assert inserted.style.name == "List Bullet"  # copied from the anchor


def test_insert_paragraph_with_explicit_style(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Beta line of content.")
    entry = {
        "id": "CL-iexp",
        "operation": "insert_paragraph_after",
        "anchor": {"paragraph_key": key},
        "new_text": "A new sub-heading.",
        "style": "Heading 2",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-iexp"] == "verified"
    edited = Document(str(path))
    inserted = next(p for p in edited.paragraphs if p.text == "A new sub-heading.")
    assert inserted.style.name == "Heading 2"


# --------------------------------------------------------------------------
# delete_paragraph
# --------------------------------------------------------------------------


def test_delete_paragraph(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Beta line of content.")
    before = texts(path)
    entry = {
        "id": "CL-del",
        "operation": "delete_paragraph",
        "anchor": {"paragraph_key": key},
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-del"] == "verified"
    after = texts(path)
    assert "Beta line of content." not in after
    assert len(after) == len(before) - 1


# --------------------------------------------------------------------------
# set_paragraph_style
# --------------------------------------------------------------------------


def test_set_paragraph_style(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {
        "id": "CL-style",
        "operation": "set_paragraph_style",
        "anchor": {"paragraph_key": key},
        "style": "Heading 2",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-style"] == "verified"
    edited = Document(str(path))
    target = next(p for p in edited.paragraphs if p.text == "Alpha line of content.")
    assert target.style.name == "Heading 2"


def test_set_paragraph_style_unknown_style_failed(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {
        "id": "CL-badstyle",
        "operation": "set_paragraph_style",
        "anchor": {"paragraph_key": key},
        "style": "No Such Style 9000",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-badstyle"] == "failed_to_apply"


# --------------------------------------------------------------------------
# replace_bullet_list
# --------------------------------------------------------------------------


def test_replace_bullet_list_shrink(tmp_path):
    path = build_bulleted(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "First bullet item.")
    entry = {
        "id": "CL-bl",
        "operation": "replace_bullet_list",
        "anchor": {"paragraph_key": key},
        "new_items": ["Replacement bullet one.", "Replacement bullet two."],
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-bl"] == "verified"
    body = texts(path)
    assert "Replacement bullet one." in body
    assert "Replacement bullet two." in body
    # The third bullet was dropped; the trailing non-list paragraph survives.
    assert "Third bullet item." not in body
    assert "Trailing non-list paragraph." in body


def test_replace_bullet_list_grow_preserves_list_style(tmp_path):
    path = build_bulleted(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "First bullet item.")
    new_items = ["b-one", "b-two", "b-three", "b-four"]
    entry = {
        "id": "CL-grow",
        "operation": "replace_bullet_list",
        "anchor": {"paragraph_key": key},
        "new_items": new_items,
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-grow"] == "verified"
    edited = Document(str(path))
    for item in new_items:
        para = next(p for p in edited.paragraphs if p.text == item)
        assert para.style.name == "List Bullet"
    assert "Trailing non-list paragraph." in [p.text for p in edited.paragraphs]


def test_replace_bullet_list_on_non_list_anchor_failed(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")  # not a list item
    entry = {
        "id": "CL-notlist",
        "operation": "replace_bullet_list",
        "anchor": {"paragraph_key": key},
        "new_items": ["x", "y"],
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-notlist"] == "failed_to_apply"
    assert "Alpha line of content." in texts(path)


# --------------------------------------------------------------------------
# operation on the versioned sample_cv fixture
# --------------------------------------------------------------------------


def test_replace_run_text_on_sample_cv_fixture(fixture_path, tmp_path):
    src = fixture_path("sample_cv.docx")
    work = tmp_path / "cv.working.docx"
    shutil.copyfile(src, work)
    doc = Document(str(work))
    bullet = "Reduced API latency by 40 percent through query and caching improvements."
    key = key_for_text(doc, bullet)
    entry = {
        "id": "CL-cv",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "40 percent"},
        "new_text": "45 percent",
    }
    result = de.apply_change_list(work, [entry], write_result=False)
    assert status_by_id(result)["CL-cv"] in {"verified", "formatting_normalized"}
    assert any("45 percent" in t for t in texts(work))


# --------------------------------------------------------------------------
# idempotency: re-running the same Change_List
# --------------------------------------------------------------------------


def test_idempotent_re_run_marks_already_satisfied(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Beta line of content.")
    change_list = [
        {
            "id": "CL-1",
            "operation": "replace_run_text",
            "anchor": {"paragraph_key": key, "match_text": "Beta"},
            "new_text": "Omega",
        }
    ]
    first = de.apply_change_list(path, change_list, write_result=False)
    assert status_by_id(first)["CL-1"] == "verified"
    body_after_first = texts(path)

    # Re-apply the exact same Change_List against the now-edited document.
    second = de.apply_change_list(path, change_list, write_result=False)
    assert status_by_id(second)["CL-1"] == "already_satisfied"
    assert second["entries"][0]["applied"] is True  # counts as applied
    # The document is unchanged by the second run (no corruption / duplication).
    assert texts(path) == body_after_first


def test_idempotent_delete_second_run_already_satisfied(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Gamma line of content.")
    change_list = [
        {"id": "CL-del", "operation": "delete_paragraph", "anchor": {"paragraph_key": key}}
    ]
    first = de.apply_change_list(path, change_list, write_result=False)
    assert status_by_id(first)["CL-del"] == "verified"
    body = texts(path)

    second = de.apply_change_list(path, change_list, write_result=False)
    assert status_by_id(second)["CL-del"] == "already_satisfied"
    assert texts(path) == body


# --------------------------------------------------------------------------
# failed_to_apply when anchor unresolved — no OTHER paragraph changes
# --------------------------------------------------------------------------


def test_failed_to_apply_unresolved_anchor_touches_nothing(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    before = texts(path)
    bogus = dn.make_paragraph_key("No Such Section", "content that does not exist", 0)
    entry = {
        "id": "CL-miss",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": bogus, "match_text": "anything"},
        "new_text": "nope",
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-miss"] == "failed_to_apply"
    assert result["failed_count"] == 1
    # Property 9: not a single other paragraph changed.
    assert texts(path) == before


def test_failed_entry_does_not_block_valid_entry(tmp_path):
    """A failing entry leaves the doc otherwise edited only where intended."""
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    valid_key = key_for_text(doc, "Alpha line of content.")
    bogus = dn.make_paragraph_key("Ghost", "ghost paragraph", 0)
    change_list = [
        {
            "id": "CL-ok",
            "operation": "replace_run_text",
            "anchor": {"paragraph_key": valid_key, "match_text": "Alpha"},
            "new_text": "Alpha-EDITED",
        },
        {
            "id": "CL-bad",
            "operation": "replace_paragraph_text",
            "anchor": {"paragraph_key": bogus},
            "new_text": "should never appear",
        },
    ]
    result = de.apply_change_list(path, change_list, write_result=False)
    statuses = status_by_id(result)
    assert statuses["CL-ok"] == "verified"
    assert statuses["CL-bad"] == "failed_to_apply"
    body = texts(path)
    assert "Alpha-EDITED line of content." in body
    assert "should never appear" not in body
    # Beta and Gamma untouched.
    assert "Beta line of content." in body
    assert "Gamma line of content." in body


# --------------------------------------------------------------------------
# verification output correctness + result.json on disk
# --------------------------------------------------------------------------


def test_result_json_written_and_statuses_match(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    alpha = key_for_text(doc, "Alpha line of content.")
    beta = key_for_text(doc, "Beta line of content.")
    bogus = dn.make_paragraph_key("Ghost", "ghost", 0)
    change_list = {
        "iteration": 3,
        "entries": [
            {
                "id": "E-verified",
                "operation": "replace_run_text",
                "anchor": {"paragraph_key": alpha, "match_text": "Alpha"},
                "new_text": "Zeta",
            },
            {
                "id": "E-deleted",
                "operation": "delete_paragraph",
                "anchor": {"paragraph_key": beta},
            },
            {
                "id": "E-failed",
                "operation": "replace_run_text",
                "anchor": {"paragraph_key": bogus, "match_text": "x"},
                "new_text": "y",
            },
        ],
    }
    cl_path = tmp_path / "iteration-3.json"
    cl_path.write_text(json.dumps(change_list), encoding="utf-8")

    result = de.apply_change_list(path, cl_path, write_result=True)

    # result.json lands next to the change list with the design's naming.
    out = tmp_path / "iteration-3.result.json"
    assert out.exists()
    on_disk = json.loads(out.read_text(encoding="utf-8"))
    assert on_disk["schema"] == de.RESULT_SCHEMA
    assert on_disk["iteration"] == 3

    statuses = {e["id"]: e["status"] for e in on_disk["entries"]}
    assert statuses == {
        "E-verified": "verified",
        "E-deleted": "verified",
        "E-failed": "failed_to_apply",
    }
    assert on_disk["applied_count"] == 2
    assert on_disk["failed_count"] == 1
    assert on_disk["counts"]["verified"] == 2
    assert on_disk["counts"]["failed_to_apply"] == 1

    # The recorded statuses match the actual on-disk document state.
    body = texts(path)
    assert "Zeta line of content." in body          # E-verified
    assert "Beta line of content." not in body       # E-deleted
    assert result["target_document"] == str(path)


def test_implements_findings_preserved_in_result(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {
        "id": "CL-trace",
        "operation": "replace_run_text",
        "anchor": {"paragraph_key": key, "match_text": "Alpha"},
        "new_text": "Trace",
        "implements_findings": ["SF-003", "ATS-011"],
    }
    result = de.apply_change_list(path, [entry], write_result=False)
    assert result["entries"][0]["implements_findings"] == ["SF-003", "ATS-011"]


# --------------------------------------------------------------------------
# Change_List loading + helpers
# --------------------------------------------------------------------------


def test_normalize_change_list_accepts_list_and_dict():
    entries, iteration = de.normalize_change_list([{"operation": "delete_paragraph"}])
    assert iteration is None and len(entries) == 1
    entries, iteration = de.normalize_change_list(
        {"iteration": 5, "entries": [{"operation": "delete_paragraph"}]}
    )
    assert iteration == 5 and len(entries) == 1


def test_normalize_change_list_rejects_bad_shapes():
    with pytest.raises(de.EditError):
        de.normalize_change_list(42)
    with pytest.raises(de.EditError):
        de.normalize_change_list({"no_entries": True})
    with pytest.raises(de.EditError):
        de.normalize_change_list(["not-an-object"])


@pytest.mark.parametrize(
    "name, expected",
    [
        ("iteration-1.json", "iteration-1.result.json"),
        ("change_list.json", "change_list.result.json"),
        ("plain", "plain.result.json"),
    ],
)
def test_result_path_for(name, expected, tmp_path):
    assert de.result_path_for(tmp_path / name) == tmp_path / expected


def test_unknown_operation_is_failed(tmp_path):
    path = build_simple(tmp_path / "cv.docx")
    doc = Document(str(path))
    key = key_for_text(doc, "Alpha line of content.")
    entry = {"id": "CL-?", "operation": "teleport_paragraph", "anchor": {"paragraph_key": key}}
    result = de.apply_change_list(path, [entry], write_result=False)
    assert status_by_id(result)["CL-?"] == "failed_to_apply"
    assert texts(path) == texts(path)  # nothing crashed; doc still readable


# --------------------------------------------------------------------------
# CLI behaviour
# --------------------------------------------------------------------------


def test_cli_applies_change_list_and_writes_result(fixture_path, tmp_path, scripts_dir):
    work = tmp_path / "cv.working.docx"
    shutil.copyfile(fixture_path("sample_cv.docx"), work)
    doc = Document(str(work))
    key = key_for_text(doc, "Python, SQL, AWS, Docker, REST APIs, testing")
    cl = {
        "iteration": 1,
        "entries": [
            {
                "id": "CLI-1",
                "operation": "replace_run_text",
                "anchor": {"paragraph_key": key, "match_text": "Docker"},
                "new_text": "Kubernetes",
            }
        ],
    }
    cl_path = tmp_path / "iteration-1.json"
    cl_path.write_text(json.dumps(cl), encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(scripts_dir / "docx_edit.py"), str(cl_path), str(work)],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["applied_count"] == 1
    assert payload["failed_count"] == 0
    assert Path(payload["result_path"]) == tmp_path / "iteration-1.result.json"
    assert (tmp_path / "iteration-1.result.json").exists()
    assert any("Kubernetes" in t for t in texts(work))


def test_cli_missing_target_exits_nonzero(tmp_path, scripts_dir):
    cl_path = tmp_path / "cl.json"
    cl_path.write_text(json.dumps([]), encoding="utf-8")
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "docx_edit.py"),
            str(cl_path),
            str(tmp_path / "missing.docx"),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == de.EXIT_INPUT_NOT_FOUND
    assert "not found" in result.stderr.lower()


def test_cli_missing_change_list_exits_nonzero(fixture_path, tmp_path, scripts_dir):
    work = tmp_path / "cv.docx"
    shutil.copyfile(fixture_path("sample_cv.docx"), work)
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "docx_edit.py"),
            str(tmp_path / "no_such_change_list.json"),
            str(work),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == de.EXIT_INPUT_NOT_FOUND
    assert "not found" in result.stderr.lower()


def test_cli_bad_change_list_exits_nonzero(fixture_path, tmp_path, scripts_dir):
    work = tmp_path / "cv.docx"
    shutil.copyfile(fixture_path("sample_cv.docx"), work)
    cl_path = tmp_path / "bad.json"
    cl_path.write_text("{ this is not valid json", encoding="utf-8")
    result = subprocess.run(
        [sys.executable, str(scripts_dir / "docx_edit.py"), str(cl_path), str(work)],
        capture_output=True,
        text=True,
    )
    assert result.returncode == de.EXIT_BAD_CHANGE_LIST


def test_missing_python_docx_exits_nonzero_naming_package(tmp_path, scripts_dir):
    """If python-docx cannot be imported, the script exits 2 naming the package."""
    blocker = tmp_path / "blocker.py"
    blocker.write_text(
        textwrap.dedent(
            f"""
            import sys, importlib.abc

            class _Block(importlib.abc.MetaPathFinder):
                def find_spec(self, name, path, target=None):
                    if name == "docx" or name.startswith("docx."):
                        raise ModuleNotFoundError(f"No module named {{name!r}}", name=name)
                    return None

            sys.meta_path.insert(0, _Block())
            sys.path.insert(0, {str(scripts_dir)!r})
            import docx_edit  # noqa: F401  (import triggers the guarded failure)
            """
        ),
        encoding="utf-8",
    )
    result = subprocess.run([sys.executable, str(blocker)], capture_output=True, text=True)
    assert result.returncode == 2
    assert "python-docx" in result.stderr


# --------------------------------------------------------------------------
# property-based: apply then re-apply is idempotent (optional, natural)
# --------------------------------------------------------------------------


@settings(max_examples=40, deadline=None)
@given(
    words=st.lists(
        st.sampled_from(["alpha", "beta", "gamma", "delta", "epsilon", "zeta"]),
        min_size=2,
        max_size=6,
        unique=True,
    ),
    repl=st.text(alphabet="XYZQW", min_size=1, max_size=6),
)
def test_property_replace_run_text_is_idempotent(words, repl):
    """Applying a replace_run_text and re-applying it leaves the doc unchanged.

    ``words`` are unique so ``match_text`` occurs exactly once, and ``repl`` uses
    a disjoint alphabet so it never reintroduces the matched token.
    """
    sentence = " ".join(words)
    match = words[0]
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_paragraph(sentence)
        doc.save(str(path))

        key = key_for_text(doc, sentence)
        change_list = [
            {
                "id": "P-1",
                "operation": "replace_run_text",
                "anchor": {"paragraph_key": key, "match_text": match},
                "new_text": repl,
            }
        ]
        first = de.apply_change_list(path, change_list, write_result=False)
        assert first["entries"][0]["status"] in {"verified", "formatting_normalized"}
        body_after_first = [p.text for p in Document(str(path)).paragraphs]

        second = de.apply_change_list(path, change_list, write_result=False)
        assert second["entries"][0]["status"] == "already_satisfied"
        assert [p.text for p in Document(str(path)).paragraphs] == body_after_first
