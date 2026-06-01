"""Unit tests for ``ats_structural.py`` (task 6 / subtask 6.1).

Covers the deterministic ATS structural-hazard detector from the design's
"ATS Layer". Two halves:

* **Hazard fixtures produce the expected Finding category/anchor.** The versioned
  ``ats_hazards.docx`` fixture bundles a text box, a 2-column section, header
  content, and a layout table; each must surface as a ``category == "ats"``
  Finding with a sensible anchor. Hazards not present in that fixture
  (non-standard heading style, parser-hostile Unicode, an embedded image) are
  exercised with small in-memory documents built here.
* **A clean doc produces none.** The versioned ``sample_cv.docx`` (clean,
  single-column, no tables/text boxes/header content) must yield zero structural
  Findings. If it ever triggered an over-eager rule, the rule — not this test —
  would be fixed.

Every Finding dict is also validated for well-formedness (required fields
present, field domains correct, ``category == "ats"``, ``status == "open"``).
Schema-file validation lands with task 7 (downstream); here we assert field
presence/domains directly.

No environment variables are used; ``ats_structural`` is importable via the path
wiring in ``conftest.py``.
"""

from __future__ import annotations

import io
import json
import struct
import subprocess
import sys
import textwrap
import zlib
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import ats_structural as ats


# --------------------------------------------------------------------------
# Finding well-formedness helper (a local stand-in for the task-7 schema)
# --------------------------------------------------------------------------

_REQUIRED_FIELDS = {
    "id",
    "source_agent",
    "iteration",
    "target_document",
    "category",
    "severity",
    "anchor",
    "current",
    "proposed",
    "rationale",
    "status",
}
_SEVERITIES = {"low", "medium", "high", "blocking"}
_STATUSES = {"open", "applied", "verification_failed", "accepted_gap", "wont_fix"}
_TARGETS = {"CV_Working_Copy", "Letter_Working_Copy", "package_coherence"}


def assert_well_formed(finding: dict, *, target="CV_Working_Copy", iteration=0) -> None:
    """Assert a single Finding dict matches the unified Finding schema domains."""
    missing = _REQUIRED_FIELDS - set(finding)
    assert not missing, f"Finding missing required fields: {sorted(missing)}"
    assert finding["category"] == "ats"
    assert finding["status"] == "open"
    assert finding["severity"] in _SEVERITIES
    assert finding["target_document"] in _TARGETS
    assert finding["target_document"] == target
    assert finding["iteration"] == iteration
    assert finding["source_agent"]
    assert isinstance(finding["id"], str) and finding["id"].startswith("ATS-")
    # Anchor is always a structured object identifying a location.
    anchor = finding["anchor"]
    assert isinstance(anchor, dict)
    assert anchor.get("type") in {"paragraph", "section", "header", "footer", "table", "unanchored"}
    assert anchor.get("hazard"), "anchor must record its hazard sub-type"
    # Human-facing fields are present strings.
    for field in ("current", "proposed", "rationale"):
        assert isinstance(finding[field], str) and finding[field]


def by_hazard(findings, hazard):
    return [f for f in findings if f["anchor"].get("hazard") == hazard]


# --------------------------------------------------------------------------
# in-memory doc builders for hazards not in the versioned fixture
# --------------------------------------------------------------------------


def _minimal_png() -> bytes:
    """Return the bytes of a tiny but valid 1x1 PNG (for add_picture)."""
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def doc_with_nonstandard_heading() -> Document:
    """A doc whose 'heading' is a Normal paragraph promoted via an outline level.

    This is exactly the ATS hazard: it *looks* like a section header and behaves
    like one in the outline, but does not use a standard Heading/Title style, so
    parsers may miss the section boundary.
    """
    doc = Document()
    doc.add_heading("Real Title", level=0)
    fake = doc.add_paragraph("Experience")  # style "Normal"
    pPr = fake._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "0")
    pPr.append(outline)
    doc.add_paragraph("A normal body paragraph under the fake heading.")
    return doc


def doc_with_hostile_unicode() -> Document:
    """A doc containing a zero-width space and a replacement character."""
    doc = Document()
    doc.add_heading("Summary", level=1)
    # U+200B ZERO WIDTH SPACE (Cf) and U+FFFD REPLACEMENT CHARACTER.
    doc.add_paragraph("Delivered\u200bresults and clean\ufffddata.")
    return doc


def doc_with_image() -> Document:
    """A doc with an embedded raster image (DrawingML a:blip)."""
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Profile", level=1)
    doc.add_picture(io.BytesIO(_minimal_png()), width=Inches(1))
    return doc


def save(doc: Document, path: Path) -> Path:
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------
# the versioned hazard fixture: text box, multi-column, header, layout table
# --------------------------------------------------------------------------


def test_hazard_fixture_surfaces_the_four_bundled_hazards(fixture_path):
    findings = ats.detect_hazards(fixture_path("ats_hazards.docx"))

    for finding in findings:
        assert_well_formed(finding)

    hazards = {f["anchor"]["hazard"] for f in findings}
    assert "text_box" in hazards
    assert "multi_column" in hazards
    assert "header_content" in hazards
    assert "layout_table" in hazards


def test_text_box_anchor_is_paragraph_scoped_with_stable_key(fixture_path):
    findings = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    boxes = by_hazard(findings, "text_box")
    assert len(boxes) == 1
    anchor = boxes[0]["anchor"]
    assert anchor["type"] == "paragraph"
    # Reuses the docx_normalize stable paragraph_key coordinate system.
    assert "paragraph_key" in anchor and "::" in anchor["paragraph_key"]
    assert boxes[0]["severity"] == "high"


def test_multi_column_anchor_identifies_the_section(fixture_path):
    findings = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    cols = by_hazard(findings, "multi_column")
    assert len(cols) == 1
    anchor = cols[0]["anchor"]
    assert anchor["type"] == "section"
    assert anchor["section_index"] == 1
    assert anchor["columns"] == 2


def test_header_anchor_identifies_the_header_part(fixture_path):
    findings = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    headers = by_hazard(findings, "header_content")
    # The fixture's two sections share one (linked) header part: flagged once.
    assert len(headers) == 1
    anchor = headers[0]["anchor"]
    assert anchor["type"] == "header"
    assert anchor["part"].endswith("header1.xml")
    assert "jordan@example.com" in headers[0]["current"]


def test_layout_table_anchor_identifies_the_table(fixture_path):
    findings = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    tables = by_hazard(findings, "layout_table")
    assert len(tables) == 1
    anchor = tables[0]["anchor"]
    assert anchor["type"] == "table"
    assert anchor["table_index"] == 0


# --------------------------------------------------------------------------
# hazards exercised with small in-memory docs
# --------------------------------------------------------------------------


def test_nonstandard_heading_detected(tmp_path):
    path = save(doc_with_nonstandard_heading(), tmp_path / "heading.docx")
    findings = ats.detect_hazards(path)
    headings = by_hazard(findings, "nonstandard_heading")
    assert len(headings) == 1
    finding = headings[0]
    assert_well_formed(finding)
    assert finding["anchor"]["type"] == "paragraph"
    assert "paragraph_key" in finding["anchor"]
    assert "Experience" in finding["current"]


def test_standard_headings_are_not_flagged_as_nonstandard(fixture_path):
    # sample_cv.docx uses Title + Heading 1 throughout: zero heading hazards.
    findings = ats.detect_hazards(fixture_path("sample_cv.docx"))
    assert by_hazard(findings, "nonstandard_heading") == []


def test_hostile_unicode_detected(tmp_path):
    path = save(doc_with_hostile_unicode(), tmp_path / "unicode.docx")
    findings = ats.detect_hazards(path)
    uni = by_hazard(findings, "hostile_unicode")
    assert len(uni) == 1
    finding = uni[0]
    assert_well_formed(finding)
    assert finding["anchor"]["type"] == "paragraph"
    # The detector names the offending code points in its rationale.
    assert "U+200B" in finding["rationale"]
    assert "U+FFFD" in finding["rationale"]


def test_ordinary_punctuation_is_not_hostile(tmp_path):
    doc = Document()
    doc.add_heading("Summary", level=1)
    # Em dash, curly quotes, and a non-breaking space are legitimate text.
    doc.add_paragraph("Senior engineer \u2014 led \u201cstrategic\u201d work\u00a0here.")
    path = save(doc, tmp_path / "punct.docx")
    findings = ats.detect_hazards(path)
    assert by_hazard(findings, "hostile_unicode") == []


def test_image_detected(tmp_path):
    path = save(doc_with_image(), tmp_path / "image.docx")
    findings = ats.detect_hazards(path)
    images = by_hazard(findings, "image_with_text")
    assert len(images) == 1
    assert_well_formed(images[0])
    assert images[0]["anchor"]["type"] == "paragraph"


# --------------------------------------------------------------------------
# the clean document produces no structural Findings
# --------------------------------------------------------------------------


def test_clean_sample_cv_produces_no_findings(fixture_path):
    findings = ats.detect_hazards(fixture_path("sample_cv.docx"))
    assert findings == [], f"clean CV unexpectedly flagged: {findings}"


def test_clean_doc_well_formed_document_wrapper(fixture_path):
    document = ats.build_findings_document(fixture_path("sample_cv.docx"))
    assert document["schema"] == ats.FINDINGS_SCHEMA
    assert document["finding_count"] == 0
    assert document["findings"] == []
    assert document["target_document"] == "CV_Working_Copy"


# --------------------------------------------------------------------------
# tagging: target_document / iteration / source_agent propagate to Findings
# --------------------------------------------------------------------------


def test_findings_carry_requested_tags(fixture_path):
    findings = ats.detect_hazards(
        fixture_path("ats_hazards.docx"),
        target_document="Letter_Working_Copy",
        iteration=3,
        source_agent="cv-ats-reviewer",
    )
    assert findings, "expected hazards on the hazard fixture"
    for finding in findings:
        assert_well_formed(finding, target="Letter_Working_Copy", iteration=3)
        assert finding["source_agent"] == "cv-ats-reviewer"


def test_invalid_target_document_rejected(fixture_path):
    with pytest.raises(ValueError):
        ats.detect_hazards(fixture_path("sample_cv.docx"), target_document="bogus")


# --------------------------------------------------------------------------
# determinism + unique ids
# --------------------------------------------------------------------------


def test_ids_are_unique_and_deterministic(fixture_path):
    first = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    second = ats.detect_hazards(fixture_path("ats_hazards.docx"))
    ids = [f["id"] for f in first]
    assert len(ids) == len(set(ids)), "Finding ids must be unique within a run"
    # Stable across runs: same ordering and same ids.
    assert ids == [f["id"] for f in second]


def test_accepts_an_open_document_object(fixture_path):
    doc = Document(str(fixture_path("ats_hazards.docx")))
    findings = ats.detect_hazards(doc)
    assert by_hazard(findings, "text_box")


# --------------------------------------------------------------------------
# CLI behaviour
# --------------------------------------------------------------------------


def test_cli_prints_findings_document(fixture_path, scripts_dir):
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "ats_structural.py"),
            str(fixture_path("ats_hazards.docx")),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["schema"] == ats.FINDINGS_SCHEMA
    assert payload["finding_count"] == len(payload["findings"]) >= 4
    for finding in payload["findings"]:
        assert_well_formed(finding)


def test_cli_writes_out_file_and_prints_summary(fixture_path, scripts_dir, tmp_path):
    out_path = tmp_path / "findings.json"
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "ats_structural.py"),
            str(fixture_path("ats_hazards.docx")),
            "--out",
            str(out_path),
            "--target",
            "Letter_Working_Copy",
            "--iteration",
            "2",
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    summary = json.loads(result.stdout)
    assert Path(summary["findings_json"]) == out_path
    assert out_path.exists()
    document = json.loads(out_path.read_text(encoding="utf-8"))
    assert document["target_document"] == "Letter_Working_Copy"
    assert document["iteration"] == 2
    for finding in document["findings"]:
        assert_well_formed(finding, target="Letter_Working_Copy", iteration=2)


def test_cli_missing_input_exits_nonzero(scripts_dir, tmp_path):
    result = subprocess.run(
        [
            sys.executable,
            str(scripts_dir / "ats_structural.py"),
            str(tmp_path / "does_not_exist.docx"),
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 2
    assert "not found" in result.stderr.lower()


def test_missing_python_docx_exits_nonzero_naming_package(tmp_path, scripts_dir):
    """If python-docx cannot be imported, the script exits 2 naming the package."""
    blocker = tmp_path / "blocker.py"
    blocker.write_text(
        textwrap.dedent(
            f"""
            import sys, importlib.abc

            class _Block(importlib.abc.MetaPathFinder):
                def find_spec(self, name, path, target=None):
                    if name == "docx" or name.startswith("docx."):
                        raise ModuleNotFoundError(f"No module named {{name!r}}", name=name)
                    return None

            sys.meta_path.insert(0, _Block())
            sys.path.insert(0, {str(scripts_dir)!r})
            import ats_structural  # noqa: F401  (import triggers the guarded failure)
            """
        ),
        encoding="utf-8",
    )
    result = subprocess.run(
        [sys.executable, str(blocker)], capture_output=True, text=True
    )
    assert result.returncode == 2
    assert "python-docx" in result.stderr
