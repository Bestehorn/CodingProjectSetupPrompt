"""docx_normalize.py — Normalized_Text + stable anchors from a `.docx` (task 2).

This is part of the deterministic Python core of the CV Customizer Agent Suite.
It does two things, deterministically and with no reasoning component:

1. Extract a ``.docx``'s body paragraphs into a normalized Markdown surface that
   reviewer agents read (``cv.normalized.md`` / ``letter.normalized.md``).
2. Emit a companion ``*.anchors.json`` sidecar mapping a **stable paragraph key**
   to the paragraph's current index in ``document.paragraphs``.

The anchor model (see design "The anchor model")
------------------------------------------------
``python-docx`` addresses paragraphs by integer index, which is unstable across
edits: inserting or deleting one paragraph shifts every later index. To make
Findings and edits robust, each paragraph is assigned a **stable key** derived
from:

  (a) the nearest preceding heading text (the paragraph's *section*),
  (b) a content hash of the paragraph's runs, and
  (c) an occurrence ordinal scoped to ``(section, content)`` that disambiguates
      identical (duplicate) paragraphs.

Because the key never encodes an absolute index, inserting or deleting *other*
paragraphs leaves a paragraph's key unchanged; only changing that paragraph's
own text (or the heading above it) invalidates its key. The editor re-derives
the mapping against the live document at edit time, so an anchor resolves
correctly even after earlier edits in the same pass. When a key resolves to
zero paragraphs (content changed underneath it), the caller reports the entry
unresolved rather than editing the wrong paragraph.

Headings anchor themselves: a heading paragraph's section is its own text, not
the heading above it. This keeps a heading's key stable even when an entire
earlier section is deleted.

CLI
---
    python docx_normalize.py <input.docx> <output.md> [anchors.json]

Given ``cv.normalized.md`` the sidecar defaults to ``cv.anchors.json`` (matching
the design's ``inputs/cv.normalized.md`` + ``cv.anchors.json`` convention). A
third positional argument overrides the sidecar path explicitly.

Dependency policy: if ``python-docx`` is not installed the script exits
non-zero naming the package. It never attempts to install anything.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ModuleNotFoundError as exc:  # pragma: no cover - exercised only when dep missing
    _missing = getattr(exc, "name", None) or "python-docx"
    sys.stderr.write(
        f"ERROR: required package not installed: {_missing}. "
        "Install 'python-docx' in this Python environment and re-run. "
        "This script does not install packages.\n"
    )
    raise SystemExit(2)


# --- constants -------------------------------------------------------------

TITLE_STYLE = "Title"
HEADING_STYLE_PREFIX = "Heading"
PREAMBLE_TOKEN = "_preamble_"
ANCHORS_SCHEMA = "docx-anchors/v1"
_CONTENT_HASH_LEN = 16
_SECTION_HASH_LEN = 8


# --- errors ----------------------------------------------------------------


class AnchorResolutionError(KeyError):
    """Raised when a paragraph_key does not resolve to exactly one paragraph."""


class AnchorCollisionError(ValueError):
    """Raised when two paragraphs produce the same key (should never happen)."""


# --- data model ------------------------------------------------------------


@dataclass(frozen=True)
class ParagraphAnchor:
    """A single body paragraph's computed anchor metadata."""

    index: int
    key: str
    section: Optional[str]
    style: str
    text: str
    is_heading: bool


# --- key construction ------------------------------------------------------


def _style_name(paragraph) -> str:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", None)
    return name or ""


def _is_heading(paragraph) -> bool:
    name = _style_name(paragraph)
    return name == TITLE_STYLE or name.startswith(HEADING_STYLE_PREFIX)


def _heading_level(style_name: str) -> int:
    """Markdown heading level offset; Title -> 0, 'Heading N' -> N, else 1."""
    if style_name == TITLE_STYLE:
        return 0
    match = re.match(r"Heading\s+(\d+)", style_name)
    return int(match.group(1)) if match else 1


def _slug(text: str, maxlen: int = 40) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.strip().lower()).strip("-")
    return slug[:maxlen].strip("-")


def _short_hash(text: str, length: int) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:length]


def _section_token(section_text: Optional[str]) -> str:
    """A readable-but-unique token for the nearest preceding heading.

    The slug keeps keys debuggable; the trailing hash of the *full* heading text
    keeps two headings that slug identically (or to nothing) in separate
    namespaces.
    """
    if section_text is None:
        return PREAMBLE_TOKEN
    slug = _slug(section_text) or "section"
    return f"{slug}-{_short_hash(section_text, _SECTION_HASH_LEN)}"


def make_paragraph_key(section_text: Optional[str], content_text: str, ordinal: int) -> str:
    """Build the stable key from (section, content hash, occurrence ordinal)."""
    return f"{_section_token(section_text)}::{_short_hash(content_text, _CONTENT_HASH_LEN)}::{ordinal}"


# --- core walk -------------------------------------------------------------


def compute_paragraph_anchors(doc) -> list[ParagraphAnchor]:
    """Walk ``doc.paragraphs`` and compute a stable anchor for each.

    The ordinal is scoped to ``(section, content)`` so identical paragraphs get
    distinct ordinals while a paragraph's ordinal is unaffected by inserting or
    deleting *other* (non-identical) paragraphs.
    """
    anchors: list[ParagraphAnchor] = []
    current_section: Optional[str] = None
    counts: dict[tuple[Optional[str], str], int] = {}

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        style_name = _style_name(paragraph)
        heading = _is_heading(paragraph)
        if heading:
            # A heading defines its own section, so it anchors to itself and
            # stays stable even if an entire earlier section is removed.
            current_section = text
        section = current_section
        group = (section, text)
        ordinal = counts.get(group, 0)
        counts[group] = ordinal + 1
        anchors.append(
            ParagraphAnchor(
                index=index,
                key=make_paragraph_key(section, text, ordinal),
                section=section,
                style=style_name,
                text=text,
                is_heading=heading,
            )
        )
    return anchors


def compute_anchor_map(doc) -> dict[str, int]:
    """Return the canonical ``paragraph_key -> current index`` mapping."""
    mapping: dict[str, int] = {}
    for anchor in compute_paragraph_anchors(doc):
        if anchor.key in mapping:
            raise AnchorCollisionError(
                f"duplicate paragraph_key {anchor.key!r} at indices "
                f"{mapping[anchor.key]} and {anchor.index}"
            )
        mapping[anchor.key] = anchor.index
    return mapping


# --- resolution ------------------------------------------------------------


def resolve_paragraph_index(doc, key: str) -> Optional[int]:
    """Resolve ``key`` against a (possibly edited) live doc.

    Returns the single matching index, or ``None`` when the key matches zero
    paragraphs (content changed underneath it) — or, impossibly, more than one.
    """
    matches = [anchor.index for anchor in compute_paragraph_anchors(doc) if anchor.key == key]
    if len(matches) == 1:
        return matches[0]
    return None


def resolve_paragraph_key(doc, key: str):
    """Resolve ``key`` to exactly one ``python-docx`` paragraph.

    Raises :class:`AnchorResolutionError` when the key does not resolve to
    exactly one paragraph, so callers never silently edit the wrong paragraph.
    """
    index = resolve_paragraph_index(doc, key)
    if index is None:
        raise AnchorResolutionError(
            f"paragraph_key {key!r} did not resolve to exactly one paragraph"
        )
    return doc.paragraphs[index]


# --- markdown rendering ----------------------------------------------------


def _markdown_line(anchor: ParagraphAnchor) -> str:
    text = anchor.text
    style = anchor.style
    if anchor.is_heading:
        hashes = "#" * (_heading_level(style) + 1)
        return f"{hashes} {text}".rstrip()
    if style.startswith("List Bullet"):
        return f"- {text}"
    if style.startswith("List Number"):
        return f"1. {text}"
    return text


def to_markdown(doc) -> str:
    """Render a deterministic normalized Markdown surface for reviewers."""
    lines: list[str] = []
    for anchor in compute_paragraph_anchors(doc):
        line = _markdown_line(anchor)
        if anchor.is_heading and lines:
            lines.append("")
        lines.append(line)
        if anchor.is_heading:
            lines.append("")
    return "\n".join(lines).strip() + "\n"


# --- anchors sidecar -------------------------------------------------------


def build_anchors_document(
    doc,
    *,
    source_document: Optional[str] = None,
    normalized_md: Optional[str] = None,
) -> dict:
    """Build the JSON-serializable anchors sidecar content.

    ``anchors`` is the canonical ``paragraph_key -> index`` mapping the editor
    re-resolves against the live document. ``paragraphs`` carries ordered
    metadata (section, style, text) for auditing and round-trip checks.
    """
    anchors = compute_paragraph_anchors(doc)
    mapping: dict[str, int] = {}
    for anchor in anchors:
        if anchor.key in mapping:
            raise AnchorCollisionError(
                f"duplicate paragraph_key {anchor.key!r} at indices "
                f"{mapping[anchor.key]} and {anchor.index}"
            )
        mapping[anchor.key] = anchor.index
    return {
        "schema": ANCHORS_SCHEMA,
        "source_document": source_document,
        "normalized_md": normalized_md,
        "paragraph_count": len(anchors),
        "anchors": mapping,
        "paragraphs": [
            {
                "index": anchor.index,
                "key": anchor.key,
                "section": anchor.section,
                "style": anchor.style,
                "is_heading": anchor.is_heading,
                "text": anchor.text,
            }
            for anchor in anchors
        ],
    }


def load_anchors(path) -> dict:
    """Load an anchors sidecar JSON file written by :func:`normalize_docx`."""
    return json.loads(Path(path).read_text(encoding="utf-8"))


def anchors_path_for(output_md: Path) -> Path:
    """Derive the sidecar path from the normalized-md output path.

    ``cv.normalized.md`` -> ``cv.anchors.json`` (design convention);
    ``foo.md`` -> ``foo.anchors.json``; ``foo`` -> ``foo.anchors.json``.
    """
    output_md = Path(output_md)
    name = output_md.name
    if name.endswith(".normalized.md"):
        stem = name[: -len(".normalized.md")]
    elif name.endswith(".md"):
        stem = name[: -len(".md")]
    else:
        stem = name
    return output_md.with_name(f"{stem}.anchors.json")


def normalize_docx(input_path, output_md_path, anchors_path=None) -> tuple[Path, Path]:
    """Normalize a ``.docx`` to Markdown + an anchors sidecar.

    Returns the ``(output_md_path, anchors_path)`` actually written.
    """
    input_path = Path(input_path)
    output_md_path = Path(output_md_path)
    anchors_path = Path(anchors_path) if anchors_path is not None else anchors_path_for(output_md_path)

    doc = Document(str(input_path))

    output_md_path.parent.mkdir(parents=True, exist_ok=True)
    output_md_path.write_text(to_markdown(doc), encoding="utf-8")

    anchors_doc = build_anchors_document(
        doc,
        source_document=input_path.name,
        normalized_md=output_md_path.name,
    )
    anchors_path.parent.mkdir(parents=True, exist_ok=True)
    anchors_path.write_text(
        json.dumps(anchors_doc, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return output_md_path, anchors_path


# --- CLI -------------------------------------------------------------------


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        prog="docx_normalize.py",
        description="Normalize a .docx to Markdown plus a stable-anchor sidecar.",
    )
    parser.add_argument("input_docx", help="Path to the input .docx file.")
    parser.add_argument("output_md", help="Path to write the normalized Markdown.")
    parser.add_argument(
        "anchors_json",
        nargs="?",
        default=None,
        help="Optional explicit anchors sidecar path (defaults next to output_md).",
    )
    args = parser.parse_args(argv)

    input_path = Path(args.input_docx)
    if not input_path.exists():
        sys.stderr.write(f"ERROR: input .docx not found: {input_path}\n")
        return 2

    out_md, out_anchors = normalize_docx(input_path, args.output_md, args.anchors_json)
    print(
        json.dumps(
            {
                "input_docx": str(input_path),
                "normalized_md": str(out_md),
                "anchors_json": str(out_anchors),
            }
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
