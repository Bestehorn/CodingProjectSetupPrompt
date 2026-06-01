"""docx_edit.py — the deterministic edit engine (task 5).

Part of the deterministic Python core of the CV Customizer Agent Suite. The CV
Editor Agent writes a thin wrapper that points this engine at a Change_List JSON
and a target ``.docx``; the engine applies a *closed* set of operations, writes
the edited document in place, and emits a per-entry verification result
(``iteration-<n>.result.json``).

Why a closed operation set
--------------------------
Keeping the vocabulary small and explicit makes the editor deterministic and
testable, and keeps the LLM's job at edit time to "fill in parameters", not
"invent docx manipulation" (design "Edit-Application Layer"). The operations are:

``replace_run_text``
    Replace a substring (``anchor.match_text``) within a paragraph's runs with
    ``new_text``, **preserving run formatting** when the match lies fully within
    a single run. When the match spans multiple runs, the affected runs are
    flattened into one run that inherits the *first* matched run's formatting and
    the entry is flagged ``formatting_normalized`` so a later language pass can
    catch any unwanted flattening.

``replace_paragraph_text``
    Replace an entire paragraph's text with ``new_text`` (keeps the first run's
    formatting and the paragraph style; drops other run-level formatting).

``insert_paragraph_after`` / ``insert_paragraph_before``
    Add a paragraph with ``new_text`` relative to the anchor, copying the
    anchor's paragraph style unless the entry supplies an explicit ``style``.

``delete_paragraph``
    Remove the anchored paragraph (used by length reduction).

``set_paragraph_style``
    Change the anchored paragraph's named style to ``style`` (e.g. fix an
    orphaned heading).

``replace_bullet_list``
    Replace the contiguous run of list items that begins at the anchor with
    ``new_items`` (a list of strings), preserving the first item's list style.

Change_List entry shape (design "Change_List entry")
----------------------------------------------------
::

    {
      "id": "CL-1-007",                       # stable id within the run
      "iteration": 1,
      "target_document": "CV_Working_Copy",   # informational; caller routes files
      "implements_findings": ["SF-003"],      # backreference(s) to Finding ids
      "operation": "replace_run_text",
      "anchor": { "paragraph_key": "...", "match_text": "QuickSuite" },
      "new_text": "Amazon Q",                 # replace_*/insert_* operations
      "style": "Heading 2",                   # set_paragraph_style / optional insert style
      "new_items": ["...", "..."],            # replace_bullet_list only
      "notes": "..."                          # free text, ignored by the engine
    }

A Change_List file is either a JSON list of entries or a dict with an
``entries`` key (and optional ``iteration``). The engine applies *every* entry
in the list to the single target ``.docx`` it is pointed at; the orchestrator is
responsible for routing CV vs. letter entries to separate invocations.

Anchor model & idempotency
---------------------------
Anchors are the stable ``paragraph_key`` from ``docx_normalize.py``. The engine
**re-resolves every anchor against the live document** immediately before
applying its entry, so multiple operations in one Change_List stay correct as
the document changes. When an anchor does not resolve to exactly one paragraph
the entry is marked ``failed_to_apply`` and **no other paragraph is touched**
(design Property 9). Before a ``replace_run_text`` edit, if ``new_text`` is
already present and ``match_text`` absent, the entry is marked
``already_satisfied`` instead of re-editing (the first line of defence against
oscillation). The other operations carry analogous post-condition checks so a
repeated Change_List is idempotent.

Verification
------------
After applying all entries the engine **saves the document, re-reads it from
disk**, and recomputes the text/style at each edited location, recording a
per-entry status drawn from ``verified | failed_to_apply | already_satisfied |
formatting_normalized``.

CLI
---
::

    python docx_edit.py <change_list.json> <target.docx> [--result <result.json>]
                        [--iteration N]

The result path defaults next to the Change_List: ``iteration-1.json`` ->
``iteration-1.result.json``. Per-entry ``failed_to_apply`` is a recorded result,
not a process error, so the CLI still exits 0; it exits non-zero only for a
missing input file, an unreadable Change_List, or a missing dependency.

Dependency policy: if ``python-docx`` is not installed the script exits non-zero
naming the package. It never attempts to install anything.
"""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
except ModuleNotFoundError as exc:  # pragma: no cover - exercised only when dep missing
    _missing = getattr(exc, "name", None) or "python-docx"
    sys.stderr.write(
        f"ERROR: required package not installed: {_missing}. "
        "Install 'python-docx' in this Python environment and re-run. "
        "This script does not install packages.\n"
    )
    raise SystemExit(2)

# docx_normalize lives alongside this script (added to sys.path by the script
# directory at CLI time, and by conftest.py at test time). python-docx is known
# importable here, so importing it will not trip its own dependency guard.
import docx_normalize as dn


__all__ = [
    "OPERATIONS",
    "RESULT_SCHEMA",
    "EditError",
    "apply_change_list",
    "load_change_list",
    "normalize_change_list",
    "result_path_for",
]


# --- constants -------------------------------------------------------------

RESULT_SCHEMA = "docx-edit-result/v1"

OP_REPLACE_RUN_TEXT = "replace_run_text"
OP_REPLACE_PARAGRAPH_TEXT = "replace_paragraph_text"
OP_INSERT_AFTER = "insert_paragraph_after"
OP_INSERT_BEFORE = "insert_paragraph_before"
OP_DELETE_PARAGRAPH = "delete_paragraph"
OP_SET_PARAGRAPH_STYLE = "set_paragraph_style"
OP_REPLACE_BULLET_LIST = "replace_bullet_list"

OPERATIONS = frozenset(
    {
        OP_REPLACE_RUN_TEXT,
        OP_REPLACE_PARAGRAPH_TEXT,
        OP_INSERT_AFTER,
        OP_INSERT_BEFORE,
        OP_DELETE_PARAGRAPH,
        OP_SET_PARAGRAPH_STYLE,
        OP_REPLACE_BULLET_LIST,
    }
)

# Per-entry status domain (design "Verification").
STATUS_VERIFIED = "verified"
STATUS_FAILED = "failed_to_apply"
STATUS_ALREADY_SATISFIED = "already_satisfied"
STATUS_FORMATTING_NORMALIZED = "formatting_normalized"

_APPLIED_STATUSES = frozenset(
    {STATUS_VERIFIED, STATUS_ALREADY_SATISFIED, STATUS_FORMATTING_NORMALIZED}
)

_LIST_STYLE_PREFIX = "List"

# Exit codes (CLI).
EXIT_OK = 0
EXIT_INPUT_NOT_FOUND = 2
EXIT_BAD_CHANGE_LIST = 3


class EditError(Exception):
    """Raised for malformed Change_List input the engine cannot apply at all."""


# --- apply-phase bookkeeping ----------------------------------------------

# Internal apply-phase stages (resolved to a public status during verification).
_STAGE_FAILED = "failed"
_STAGE_ALREADY = "already"
_STAGE_APPLIED = "applied"

# verify kinds
_VK_TEXT = "text"
_VK_STYLE = "style"
_VK_DELETE = "delete"
_VK_BULLET = "bullet"


@dataclass
class _Pending:
    """Apply-phase outcome for a single Change_List entry, verified later."""

    entry_id: Optional[str]
    operation: str
    implements_findings: list = field(default_factory=list)
    stage: str = _STAGE_FAILED
    reason: Optional[str] = None
    normalized: bool = False
    verify_kind: Optional[str] = None
    elem: Any = None
    expected_text: Optional[str] = None
    expected_style: Optional[str] = None
    expected_items: Optional[list] = None
    anchor_key: Optional[str] = None


# --- small docx helpers ----------------------------------------------------


def _runs_text(paragraph) -> str:
    """Concatenated text of a paragraph's direct runs (offset-stable surface)."""
    return "".join(run.text or "" for run in paragraph.runs)


def _is_list_style(paragraph) -> bool:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", None) or ""
    return name.startswith(_LIST_STYLE_PREFIX)


def _remove_run(run) -> None:
    element = run._element
    element.getparent().remove(element)


def _delete_paragraph_element(paragraph) -> None:
    element = paragraph._p
    element.getparent().remove(element)


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's text, keeping the first run's char formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            _remove_run(run)
    else:
        paragraph.add_run(new_text)


def _insert_paragraph(anchor_para, text: str, style, *, before: bool):
    """Insert a new paragraph adjacent to ``anchor_para`` with ``text``.

    Copies ``style`` (a style object or name) onto the new paragraph; ``style``
    of ``None`` leaves the default. Returns the new :class:`Paragraph`.
    """
    new_p = OxmlElement("w:p")
    if before:
        anchor_para._p.addprevious(new_p)
    else:
        anchor_para._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor_para._parent)
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def _section_token_of(key: str) -> str:
    """The stable section component of a ``paragraph_key`` (before the first ``::``)."""
    return key.split("::", 1)[0]


@dataclass(frozen=True)
class _LiveRow:
    index: int
    section_token: str
    text: str
    is_list: bool


def _live_rows(doc) -> list:
    """Snapshot the live doc as ``(index, section_token, text, is_list)`` rows.

    Used by the idempotency guards when an edit has already invalidated its own
    content-hash anchor key (a re-run of the same Change_List): the section token
    is derived from the unchanged nearest-heading, so we can still recognise the
    post-edit state within the right section without editing anything.
    """
    rows = []
    for anchor in dn.compute_paragraph_anchors(doc):
        rows.append(
            _LiveRow(
                index=anchor.index,
                section_token=_section_token_of(anchor.key),
                text=anchor.text,
                is_list=anchor.style.startswith(_LIST_STYLE_PREFIX),
            )
        )
    return rows


def _bullet_block_present(rows, section_token: str, new_items: list) -> bool:
    """True if ``new_items`` appears as a contiguous list block in the section."""
    n = len(new_items)
    if n == 0:
        return False
    i = 0
    while i < len(rows):
        if rows[i].is_list and rows[i].section_token == section_token:
            j = i
            block = []
            while j < len(rows) and rows[j].is_list:
                block.append(rows[j].text)
                j += 1
            if block == new_items:
                return True
            i = j
        else:
            i += 1
    return False


def _replace_run_text(paragraph, match_text: str, new_text: str) -> bool:
    """Replace the first occurrence of ``match_text`` within a paragraph.

    Returns ``True`` when the match spanned multiple runs and was flattened into
    a single run (``formatting_normalized``); ``False`` when the edit stayed
    fully within one run (formatting preserved in place).
    """
    runs = paragraph.runs
    full = "".join(run.text or "" for run in runs)
    start = full.find(match_text)
    end = start + len(match_text)

    spans = []
    pos = 0
    for run in runs:
        length = len(run.text or "")
        spans.append((run, pos, pos + length))
        pos += length

    affected = [(run, s, e) for (run, s, e) in spans if e > start and s < end]

    if len(affected) == 1:
        run, s, _e = affected[0]
        local_start = start - s
        local_end = end - s
        original = run.text or ""
        run.text = original[:local_start] + new_text + original[local_end:]
        return False

    first_run, fs, _fe = affected[0]
    last_run, ls, _le = affected[-1]
    prefix = (first_run.text or "")[: start - fs]
    suffix = (last_run.text or "")[end - ls :]
    first_run.text = prefix + new_text + suffix
    for run, _s, _e in affected[1:]:
        _remove_run(run)
    return True


# --- per-entry apply -------------------------------------------------------


def _fail(pending: _Pending, reason: str) -> _Pending:
    pending.stage = _STAGE_FAILED
    pending.reason = reason
    return pending


def _already_satisfied_when_unresolved(doc, pending: _Pending, entry: dict, anchor: dict, key: str) -> _Pending:
    """Decide an unresolved-anchor entry: ``already_satisfied`` vs ``failed``.

    A content-mutating edit (``replace_run_text`` / ``replace_paragraph_text`` /
    ``replace_bullet_list``) changes its target paragraph's content-hash key, so
    re-applying the *same* Change_List finds the original anchor unresolved even
    though the intended result is already present (oscillation defence, design
    "Idempotency"). Because the anchor encodes the nearest-heading section token
    (which the edit did NOT change), we re-check the operation's post-condition
    *scoped to that section* without mutating anything. If the post-condition
    holds the entry is ``already_satisfied``; otherwise it is a genuine
    ``failed_to_apply`` (no paragraph is touched either way — Property 9).

    Insert/style operations do not change their own anchor's key, so for them an
    unresolved anchor is always a genuine failure.
    """
    operation = pending.operation
    section_token = _section_token_of(key)
    rows = _live_rows(doc)
    section_rows = [r for r in rows if r.section_token == section_token]
    fail_reason = f"anchor {key!r} did not resolve to a paragraph"

    if operation == OP_REPLACE_RUN_TEXT:
        match_text = anchor.get("match_text")
        new_text = entry.get("new_text")
        if not match_text:
            return _fail(pending, "replace_run_text requires anchor.match_text")
        if new_text is None:
            return _fail(pending, "replace_run_text requires new_text")
        for row in section_rows:
            if new_text and new_text in row.text and match_text not in row.text:
                pending.stage = _STAGE_ALREADY
                return pending
        return _fail(pending, fail_reason)

    if operation == OP_REPLACE_PARAGRAPH_TEXT:
        new_text = entry.get("new_text")
        if new_text is None:
            return _fail(pending, "replace_paragraph_text requires new_text")
        for row in section_rows:
            if row.text == new_text:
                pending.stage = _STAGE_ALREADY
                return pending
        return _fail(pending, fail_reason)

    if operation == OP_REPLACE_BULLET_LIST:
        new_items = entry.get("new_items")
        if not isinstance(new_items, list) or not all(isinstance(x, str) for x in new_items):
            return _fail(pending, "replace_bullet_list requires new_items: list[str]")
        if not new_items:
            return _fail(pending, "replace_bullet_list requires a non-empty new_items")
        if _bullet_block_present(rows, section_token, new_items):
            pending.stage = _STAGE_ALREADY
            return pending
        return _fail(pending, fail_reason)

    return _fail(pending, fail_reason)


def _apply_entry(doc, entry: dict) -> _Pending:
    """Mutate ``doc`` for one entry; record a :class:`_Pending` (no verify yet).

    A failed entry mutates nothing (Property 9): every failure path returns
    before touching the document.
    """
    operation = entry.get("operation")
    pending = _Pending(
        entry_id=entry.get("id"),
        operation=operation if isinstance(operation, str) else str(operation),
        implements_findings=list(entry.get("implements_findings") or []),
    )

    if operation not in OPERATIONS:
        return _fail(pending, f"unknown operation {operation!r}")

    anchor = entry.get("anchor") or {}
    if not isinstance(anchor, dict):
        return _fail(pending, "anchor must be an object")
    paragraph_key = anchor.get("paragraph_key")
    pending.anchor_key = paragraph_key
    if not paragraph_key:
        return _fail(pending, "anchor.paragraph_key is required")

    index = dn.resolve_paragraph_index(doc, paragraph_key)

    # delete is idempotent: an unresolved anchor means it is already gone.
    if operation == OP_DELETE_PARAGRAPH:
        if index is None:
            pending.stage = _STAGE_ALREADY
            return pending
        paragraph = doc.paragraphs[index]
        pending.elem = paragraph._p  # captured for completeness; verified via key
        pending.verify_kind = _VK_DELETE
        _delete_paragraph_element(paragraph)
        pending.stage = _STAGE_APPLIED
        return pending

    if index is None:
        return _already_satisfied_when_unresolved(doc, pending, entry, anchor, paragraph_key)
    paragraph = doc.paragraphs[index]

    if operation == OP_REPLACE_RUN_TEXT:
        match_text = anchor.get("match_text")
        new_text = entry.get("new_text")
        if not match_text:
            return _fail(pending, "replace_run_text requires anchor.match_text")
        if new_text is None:
            return _fail(pending, "replace_run_text requires new_text")
        full = _runs_text(paragraph)
        if match_text in full:
            normalized = _replace_run_text(paragraph, match_text, new_text)
            pending.normalized = normalized
            pending.verify_kind = _VK_TEXT
            pending.elem = paragraph._p
            pending.expected_text = full.replace(match_text, new_text, 1)
            pending.stage = _STAGE_APPLIED
            return pending
        # match absent: idempotency — already satisfied if the target text holds.
        if new_text == "" or (new_text and new_text in full):
            pending.stage = _STAGE_ALREADY
            return pending
        return _fail(pending, f"match_text {match_text!r} not found at anchor")

    if operation == OP_REPLACE_PARAGRAPH_TEXT:
        new_text = entry.get("new_text")
        if new_text is None:
            return _fail(pending, "replace_paragraph_text requires new_text")
        if _runs_text(paragraph) == new_text:
            pending.stage = _STAGE_ALREADY
            return pending
        _set_paragraph_text(paragraph, new_text)
        pending.verify_kind = _VK_TEXT
        pending.elem = paragraph._p
        pending.expected_text = new_text
        pending.stage = _STAGE_APPLIED
        return pending

    if operation in (OP_INSERT_AFTER, OP_INSERT_BEFORE):
        new_text = entry.get("new_text")
        if new_text is None:
            return _fail(pending, f"{operation} requires new_text")
        before = operation == OP_INSERT_BEFORE
        neighbor_index = index - 1 if before else index + 1
        if 0 <= neighbor_index < len(doc.paragraphs):
            if _runs_text(doc.paragraphs[neighbor_index]) == new_text:
                pending.stage = _STAGE_ALREADY
                return pending
        style_name = entry.get("style")
        style = style_name if style_name else paragraph.style
        new_para = _insert_paragraph(paragraph, new_text, style, before=before)
        pending.verify_kind = _VK_TEXT
        pending.elem = new_para._p
        pending.expected_text = new_text
        pending.stage = _STAGE_APPLIED
        return pending

    if operation == OP_SET_PARAGRAPH_STYLE:
        style_name = entry.get("style")
        if not style_name:
            return _fail(pending, "set_paragraph_style requires style")
        current = getattr(getattr(paragraph, "style", None), "name", None)
        if current == style_name:
            pending.stage = _STAGE_ALREADY
            return pending
        try:
            paragraph.style = style_name
        except KeyError:
            return _fail(pending, f"unknown style {style_name!r}")
        pending.verify_kind = _VK_STYLE
        pending.elem = paragraph._p
        pending.expected_style = style_name
        pending.stage = _STAGE_APPLIED
        return pending

    if operation == OP_REPLACE_BULLET_LIST:
        new_items = entry.get("new_items")
        if not isinstance(new_items, list) or not all(isinstance(x, str) for x in new_items):
            return _fail(pending, "replace_bullet_list requires new_items: list[str]")
        if not new_items:
            return _fail(pending, "replace_bullet_list requires a non-empty new_items")
        paras = doc.paragraphs
        if not _is_list_style(paras[index]):
            return _fail(pending, "anchor for replace_bullet_list is not a list item")
        block = []
        i = index
        while i < len(paras) and _is_list_style(paras[i]):
            block.append(paras[i])
            i += 1
        existing_texts = [_runs_text(p) for p in block]
        if existing_texts == new_items:
            pending.stage = _STAGE_ALREADY
            pending.verify_kind = _VK_BULLET
            pending.elem = block[0]._p
            pending.expected_items = list(new_items)
            return pending
        list_style = block[0].style
        first_elem = block[0]._p
        overlap = min(len(block), len(new_items))
        for k in range(overlap):
            _set_paragraph_text(block[k], new_items[k])
        if len(new_items) > len(block):
            prev = block[-1]
            for k in range(len(block), len(new_items)):
                prev = _insert_paragraph(prev, new_items[k], list_style, before=False)
        elif len(block) > len(new_items):
            for extra in block[len(new_items):]:
                _delete_paragraph_element(extra)
        pending.verify_kind = _VK_BULLET
        pending.elem = first_elem
        pending.expected_items = list(new_items)
        pending.stage = _STAGE_APPLIED
        return pending

    return _fail(pending, f"unhandled operation {operation!r}")  # pragma: no cover


# --- verification ----------------------------------------------------------

# Verification re-reads the saved ``.docx`` from disk, so the live element
# objects from the apply phase are gone. We bridge the two worlds by computing,
# *after all entries are applied but before saving*, the final index of each
# touched ``w:p`` element in ``doc.paragraphs``. Saving a docx preserves
# paragraph order and count, so index ``i`` in the in-memory final document is
# index ``i`` in the re-read verification document. Deletes leave no surviving
# element, so they are verified against the round-tripped paragraph count
# (which avoids the ordinal-shift pitfall of re-resolving a removed key).


@dataclass
class _VerifyContext:
    verify_doc: Any
    index_by_elem: dict
    expected_total: int
    roundtrip_ok: bool


def _verify_pending(pending: _Pending, ctx: _VerifyContext) -> str:
    """Re-read-based verification → a public per-entry status string."""
    if pending.stage == _STAGE_FAILED:
        return STATUS_FAILED
    if pending.stage == _STAGE_ALREADY:
        return STATUS_ALREADY_SATISFIED

    paragraphs = ctx.verify_doc.paragraphs
    kind = pending.verify_kind

    if kind == _VK_DELETE:
        # The paragraph is gone in-memory; a faithful round-trip (same final
        # paragraph count on disk) confirms the deletion persisted.
        ok = ctx.roundtrip_ok and len(paragraphs) == ctx.expected_total
        return STATUS_VERIFIED if ok else STATUS_FAILED

    idx = ctx.index_by_elem.get(id(pending.elem))
    if idx is None or idx >= len(paragraphs):
        return STATUS_FAILED

    if kind == _VK_TEXT:
        if _runs_text(paragraphs[idx]) != pending.expected_text:
            return STATUS_FAILED
        return STATUS_FORMATTING_NORMALIZED if pending.normalized else STATUS_VERIFIED

    if kind == _VK_STYLE:
        name = getattr(getattr(paragraphs[idx], "style", None), "name", None)
        return STATUS_VERIFIED if name == pending.expected_style else STATUS_FAILED

    if kind == _VK_BULLET:
        expected = pending.expected_items or []
        if idx + len(expected) > len(paragraphs):
            return STATUS_FAILED
        actual = [_runs_text(paragraphs[idx + k]) for k in range(len(expected))]
        return STATUS_VERIFIED if actual == expected else STATUS_FAILED

    return STATUS_FAILED  # pragma: no cover


# --- Change_List loading / normalization -----------------------------------


def normalize_change_list(payload: Any) -> tuple[list, Optional[int]]:
    """Coerce a parsed Change_List into ``(entries, iteration)``.

    Accepts either a JSON list of entries, or a dict with an ``entries`` list
    (and optional top-level ``iteration``). Raises :class:`EditError` for any
    other shape.
    """
    if isinstance(payload, list):
        entries = payload
        iteration = None
    elif isinstance(payload, dict):
        entries = payload.get("entries")
        iteration = payload.get("iteration")
        if entries is None:
            raise EditError("Change_List object must contain an 'entries' list")
    else:
        raise EditError("Change_List must be a JSON list or an object with 'entries'")
    if not isinstance(entries, list):
        raise EditError("Change_List 'entries' must be a list")
    for i, entry in enumerate(entries):
        if not isinstance(entry, dict):
            raise EditError(f"Change_List entry #{i} must be an object")
    return entries, iteration


def load_change_list(path) -> tuple[list, Optional[int]]:
    """Load and normalize a Change_List JSON file from disk."""
    raw = Path(path).read_text(encoding="utf-8")
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise EditError(f"Change_List is not valid JSON: {exc}") from exc
    return normalize_change_list(payload)


def result_path_for(change_list_path) -> Path:
    """Derive the result path: ``iteration-1.json`` -> ``iteration-1.result.json``."""
    path = Path(change_list_path)
    name = path.name
    stem = name[: -len(".json")] if name.endswith(".json") else name
    return path.with_name(f"{stem}.result.json")


# --- top-level apply + verify ----------------------------------------------


def apply_change_list(
    target_docx,
    change_list,
    *,
    iteration: Optional[int] = None,
    result_path=None,
    write_result: bool = True,
) -> dict:
    """Apply a Change_List to ``target_docx`` in place and verify the result.

    ``change_list`` may be a path to a Change_List JSON file, a parsed payload
    (list or ``{"entries": [...]}`` dict), or an already-normalized list of
    entry dicts. Every entry is applied to the single ``target_docx``; the
    document is saved, re-read from disk, and each entry verified.

    Returns the result document (also written to ``result_path`` as
    ``iteration-<n>.result.json`` when ``write_result`` is true). The result has
    a per-entry ``status`` drawn from ``verified | failed_to_apply |
    already_satisfied | formatting_normalized``.
    """
    target_docx = Path(target_docx)

    if isinstance(change_list, (str, Path)):
        entries, file_iteration = load_change_list(change_list)
        if result_path is None and write_result:
            result_path = result_path_for(change_list)
    else:
        entries, file_iteration = normalize_change_list(change_list)
    if iteration is None:
        iteration = file_iteration

    doc = Document(str(target_docx))

    # Apply every entry against the LIVE document, re-resolving each anchor just
    # before its edit so sequential operations stay correct as the doc changes.
    pendings: list[_Pending] = [_apply_entry(doc, entry) for entry in entries]

    # Final index of every surviving touched element (order is preserved on save).
    index_by_elem = {id(p._p): i for i, p in enumerate(doc.paragraphs)}
    expected_total = len(doc.paragraphs)

    # Save ONLY when at least one entry actually mutated the document. Entries
    # that resolved to ``already_satisfied`` or ``failed_to_apply`` change
    # nothing, so re-saving would needlessly rewrite the ``.docx`` zip container
    # and alter its bytes even though its content is identical — breaking the
    # idempotency contract at the file level (design "Idempotency"; the first
    # line of defence against oscillation). A no-op Change_List must leave the
    # working copy byte-for-byte unchanged.
    mutated = any(pending.stage == _STAGE_APPLIED for pending in pendings)
    if mutated:
        doc.save(str(target_docx))

    # Verification: re-read from disk and recompute text/style at each location.
    # When nothing mutated, the on-disk file is the (unchanged) original — the
    # re-read still yields a faithful document for the round-trip checks, and
    # ``already_satisfied`` / ``failed_to_apply`` entries are resolved without
    # consulting it.
    verify_doc = Document(str(target_docx))
    roundtrip_ok = len(verify_doc.paragraphs) == expected_total
    ctx = _VerifyContext(
        verify_doc=verify_doc,
        index_by_elem=index_by_elem,
        expected_total=expected_total,
        roundtrip_ok=roundtrip_ok,
    )

    entry_results = []
    counts = {
        STATUS_VERIFIED: 0,
        STATUS_FAILED: 0,
        STATUS_ALREADY_SATISFIED: 0,
        STATUS_FORMATTING_NORMALIZED: 0,
    }
    for pending in pendings:
        status = _verify_pending(pending, ctx)
        counts[status] += 1
        record = {
            "id": pending.entry_id,
            "operation": pending.operation,
            "implements_findings": pending.implements_findings,
            "status": status,
            "applied": status in _APPLIED_STATUSES,
        }
        if status == STATUS_FORMATTING_NORMALIZED:
            record["formatting_normalized"] = True
        if pending.reason:
            record["reason"] = pending.reason
        entry_results.append(record)

    result = {
        "schema": RESULT_SCHEMA,
        "target_document": str(target_docx),
        "iteration": iteration,
        "entry_count": len(entries),
        "applied_count": sum(counts[s] for s in _APPLIED_STATUSES),
        "failed_count": counts[STATUS_FAILED],
        "counts": counts,
        "entries": entry_results,
    }

    if write_result and result_path is not None:
        result_path = Path(result_path)
        result_path.parent.mkdir(parents=True, exist_ok=True)
        result_path.write_text(
            json.dumps(result, indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )
        result["result_path"] = str(result_path)

    return result


# --- CLI -------------------------------------------------------------------


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        prog="docx_edit.py",
        description="Apply a closed-vocabulary Change_List to a .docx and verify it.",
    )
    parser.add_argument("change_list", help="Path to the Change_List JSON file.")
    parser.add_argument("target_docx", help="Path to the target .docx (edited in place).")
    parser.add_argument(
        "--result",
        default=None,
        help="Where to write the result JSON (defaults to <change_list>.result.json).",
    )
    parser.add_argument(
        "--iteration",
        type=int,
        default=None,
        help="Iteration number to record in the result (overrides the Change_List's).",
    )
    args = parser.parse_args(argv)

    change_list_path = Path(args.change_list)
    target_docx = Path(args.target_docx)

    if not change_list_path.exists():
        sys.stderr.write(f"ERROR: Change_List not found: {change_list_path}\n")
        return EXIT_INPUT_NOT_FOUND
    if not target_docx.exists():
        sys.stderr.write(f"ERROR: target .docx not found: {target_docx}\n")
        return EXIT_INPUT_NOT_FOUND

    result_path = Path(args.result) if args.result else result_path_for(change_list_path)

    try:
        result = apply_change_list(
            target_docx,
            change_list_path,
            iteration=args.iteration,
            result_path=result_path,
            write_result=True,
        )
    except EditError as exc:
        sys.stderr.write(f"ERROR: {exc}\n")
        return EXIT_BAD_CHANGE_LIST

    print(
        json.dumps(
            {
                "target_document": result["target_document"],
                "result_path": result.get("result_path"),
                "applied_count": result["applied_count"],
                "failed_count": result["failed_count"],
                "counts": result["counts"],
            }
        )
    )
    return EXIT_OK


if __name__ == "__main__":
    raise SystemExit(main())
