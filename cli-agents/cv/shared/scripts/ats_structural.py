"""ats_structural.py — deterministic ATS structural-hazard detector (task 6).

Part of the deterministic Python core of the CV Customizer Agent Suite. The ATS
Reviewer Agent (`cv-ats-reviewer`) runs this via the script-in-tmp + shell
pattern to find ATS-hostile *structure* in a ``.docx`` — the objectively
detectable hazards that live in the document XML — and emits **candidate
Findings** as JSON that the agent then adopts. Keyword matching against the Job
Description is the LLM's job and lives in the agent prompt, not here.

Hazards detected (all via ``python-docx`` XML inspection)
---------------------------------------------------------
* ``text_box``          — text inside a text box (``w:txbxContent``); ATS often
                          cannot read it at all.
* ``image_with_text``   — embedded images/drawings (DrawingML ``a:blip`` or VML
                          ``v:imagedata``); any text baked into a graphic is lost.
* ``multi_column``      — section properties with ``w:cols`` ``w:num`` > 1; read
                          out of order by many parsers.
* ``header_content`` /  — non-empty header/footer that defines its own content;
  ``footer_content``      commonly ignored entirely by ATS.
* ``layout_table``      — any table (``doc.tables``); cells are frequently read
                          out of order or merged.
* ``nonstandard_heading``— a heading-like paragraph (carries an outline level)
                          that does *not* use a standard ``Title`` / ``Heading N``
                          style; section boundaries may go unrecognised.
* ``hostile_unicode``   — control / format / private-use / surrogate / unassigned
                          characters (and U+FFFD) that fail to round-trip through
                          naive parsers. Ordinary punctuation (em dash, curly
                          quotes, NBSP) is intentionally NOT flagged.

The anchor convention
---------------------
Every Finding's ``anchor`` is a JSON object carrying a ``type`` and a ``hazard``
sub-type so the orchestrator/editor know where and what:

* ``type == "paragraph"`` — the hazard maps to a single body paragraph. Carries
  the stable ``paragraph_key`` from :mod:`docx_normalize` (so the editor can act
  on it with the same coordinate system as every other Finding), plus the
  ``section`` heading and a ``text`` snippet. Used for text boxes, images,
  non-standard headings, and hostile-unicode paragraphs.
* ``type == "section"`` — a whole multi-column section. Carries ``section_index``
  and the column count.
* ``type == "header"`` / ``"footer"`` — header/footer content. Carries
  ``section_index`` and the ``part`` name (e.g. ``/word/header1.xml``).
* ``type == "table"`` — a layout table. Carries ``table_index``.
* ``type == "unanchored"`` — fallback when a body paragraph cannot be located
  (should not happen for well-formed documents).

Keeping the stable ``paragraph_key`` on paragraph-scoped hazards means a text-box
or hostile-unicode Finding can be fed straight into ``docx_edit.py`` once the
content is relocated, exactly like a spelling Finding.

The id scheme
-------------
``ATS-<TYPECODE>-<hash8>`` where ``TYPECODE`` is a short per-hazard code
(``TXB``, ``IMG``, ``COL``, ``HDR``, ``FTR``, ``TBL``, ``HDG``, ``UNI``) and
``hash8`` is the first 8 hex chars of the SHA-1 of a stable *identity* string
built from the hazard type and its location (``paragraph_key`` for
paragraph-scoped hazards; ``section_index`` / ``table_index`` / header part name
otherwise). The id is therefore deterministic and stable across runs for the
same document, and independent of detection order.
In the astronomically unlikely event two identities collide, a ``-2``/``-3``
suffix disambiguates so ids stay unique within a run.

CLI
---
::

    python ats_structural.py <docx> [--out findings.json]
        [--target CV_Working_Copy|Letter_Working_Copy] [--iteration N]
        [--source-agent NAME]

With ``--out`` the full Findings document is written there and a one-line JSON
summary is printed to stdout; without ``--out`` the full Findings document is
printed to stdout. Per design, structural Findings are *candidates*: emitting
them is never an error, so the CLI exits 0 even when hazards are found. It exits
non-zero only for a missing input file or a missing dependency.

Dependency policy: if ``python-docx`` is not installed the script exits non-zero
naming the package. It never attempts to install anything.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import unicodedata
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
except ModuleNotFoundError as exc:  # pragma: no cover - exercised only when dep missing
    _missing = getattr(exc, "name", None) or "python-docx"
    sys.stderr.write(
        f"ERROR: required package not installed: {_missing}. "
        "Install 'python-docx' in this Python environment and re-run. "
        "This script does not install packages.\n"
    )
    raise SystemExit(2)

# docx_normalize lives alongside this script (added to sys.path by the script
# directory at CLI time, and by conftest.py at test time). python-docx is known
# importable here, so importing it will not trip its own dependency guard.
import docx_normalize as dn


__all__ = [
    "CATEGORY_ATS",
    "STATUS_OPEN",
    "DEFAULT_SOURCE_AGENT",
    "VALID_TARGET_DOCUMENTS",
    "FINDINGS_SCHEMA",
    "detect_hazards",
    "build_findings_document",
    "write_findings",
]


# --- constants -------------------------------------------------------------

CATEGORY_ATS = "ats"
STATUS_OPEN = "open"
DEFAULT_SOURCE_AGENT = "cv-ats-reviewer"
FINDINGS_SCHEMA = "ats-structural-findings/v1"

TARGET_CV = "CV_Working_Copy"
TARGET_LETTER = "Letter_Working_Copy"
TARGET_PACKAGE = "package_coherence"
VALID_TARGET_DOCUMENTS = frozenset({TARGET_CV, TARGET_LETTER, TARGET_PACKAGE})

# Severity domain (design Finding field): low | medium | high | blocking.
SEV_LOW = "low"
SEV_MEDIUM = "medium"
SEV_HIGH = "high"
SEV_BLOCKING = "blocking"

# XML namespaces. python-docx's ``qn`` does not know the VML ('v') namespace, so
# we resolve Clark-notation tags ourselves for every namespace we touch.
_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
}

# Exit codes (CLI).
EXIT_OK = 0
EXIT_INPUT_NOT_FOUND = 2

_SNIPPET_LIMIT = 160


def _q(prefix: str, tag: str) -> str:
    """Return a Clark-notation ``{namespace}tag`` for ``prefix:tag``."""
    return f"{{{_NS[prefix]}}}{tag}"


def _snippet(text: Optional[str], limit: int = _SNIPPET_LIMIT) -> str:
    """Collapse whitespace and truncate ``text`` for human-readable fields."""
    collapsed = " ".join((text or "").split())
    if len(collapsed) <= limit:
        return collapsed
    return collapsed[: limit - 1] + "\u2026"


# --- hazard model ----------------------------------------------------------

# Per-hazard metadata: (typecode, severity, default proposed remediation).
_TYPE_TEXT_BOX = "text_box"
_TYPE_IMAGE = "image_with_text"
_TYPE_MULTI_COLUMN = "multi_column"
_TYPE_HEADER = "header_content"
_TYPE_FOOTER = "footer_content"
_TYPE_LAYOUT_TABLE = "layout_table"
_TYPE_NONSTANDARD_HEADING = "nonstandard_heading"
_TYPE_HOSTILE_UNICODE = "hostile_unicode"

_TYPE_CODES = {
    _TYPE_TEXT_BOX: "TXB",
    _TYPE_IMAGE: "IMG",
    _TYPE_MULTI_COLUMN: "COL",
    _TYPE_HEADER: "HDR",
    _TYPE_FOOTER: "FTR",
    _TYPE_LAYOUT_TABLE: "TBL",
    _TYPE_NONSTANDARD_HEADING: "HDG",
    _TYPE_HOSTILE_UNICODE: "UNI",
}


class _Hazard:
    """A single detected structural hazard, pre-Finding."""

    __slots__ = (
        "hazard",
        "severity",
        "anchor",
        "current",
        "proposed",
        "rationale",
        "identity",
    )

    def __init__(self, hazard, severity, anchor, current, proposed, rationale, identity):
        self.hazard = hazard
        self.severity = severity
        self.anchor = anchor
        self.current = current
        self.proposed = proposed
        self.rationale = rationale
        self.identity = identity


# --- standard-heading recognition + outline level --------------------------

_STANDARD_HEADING_RE = re.compile(r"^Heading \d+$")


def _is_standard_heading_style(style_name: str) -> bool:
    return style_name == dn.TITLE_STYLE or bool(_STANDARD_HEADING_RE.match(style_name or ""))


def _safe_int(value) -> Optional[int]:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _effective_outline_level(paragraph) -> Optional[int]:
    """Return the paragraph's effective Word outline level, or ``None``.

    Checks the paragraph's *direct* ``w:pPr/w:outlineLvl`` first (an explicit
    promotion of an otherwise-ordinary paragraph), then walks the style chain
    (``style`` → ``base_style`` → ...) for an outline level defined by the
    applied style. Word uses levels 0-8 for outline headings; 9 means body text.
    """
    pPr = paragraph._p.find(_q("w", "pPr"))
    if pPr is not None:
        lvl = pPr.find(_q("w", "outlineLvl"))
        if lvl is not None:
            return _safe_int(lvl.get(_q("w", "val")))

    style = getattr(paragraph, "style", None)
    seen: set[int] = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        element = getattr(style, "element", None)
        if element is not None:
            spPr = element.find(_q("w", "pPr"))
            if spPr is not None:
                lvl = spPr.find(_q("w", "outlineLvl"))
                if lvl is not None:
                    return _safe_int(lvl.get(_q("w", "val")))
        style = getattr(style, "base_style", None)
    return None


def _is_heading_like(paragraph) -> bool:
    level = _effective_outline_level(paragraph)
    return level is not None and 0 <= level <= 8


# --- hostile-unicode classification ----------------------------------------

# Control / format / surrogate / private-use / unassigned categories never
# survive a naive parser cleanly. Ordinary punctuation and spaces (Zs, Pd, Pi,
# Pf, etc.) are intentionally NOT hostile.
_HOSTILE_CATEGORIES = frozenset({"Cc", "Cf", "Cs", "Co", "Cn"})
# Whitespace control characters that are legitimate in document text.
_ALLOWED_CONTROL = frozenset({"\t", "\n", "\r"})
# Explicitly hostile characters that are not in a hostile *category*.
_EXPLICIT_HOSTILE = frozenset({"\ufffd"})  # U+FFFD REPLACEMENT CHARACTER


def _hostile_chars(text: str) -> list[str]:
    """Return the distinct parser-hostile characters in ``text`` (stable order)."""
    found: list[str] = []
    seen: set[str] = set()
    for ch in text or "":
        if ch in _ALLOWED_CONTROL:
            continue
        hostile = ch in _EXPLICIT_HOSTILE or unicodedata.category(ch) in _HOSTILE_CATEGORIES
        if hostile and ch not in seen:
            seen.add(ch)
            found.append(ch)
    return found


def _describe_chars(chars: list[str]) -> str:
    parts = []
    for ch in chars:
        name = unicodedata.name(ch, "UNNAMED")
        parts.append(f"U+{ord(ch):04X} ({name})")
    return ", ".join(parts)


# --- body-paragraph helpers ------------------------------------------------


def _iter_body_paragraph_anchors(doc):
    """Yield ``(paragraph, ParagraphAnchor)`` for every body paragraph.

    ``docx_normalize.compute_paragraph_anchors`` walks ``doc.paragraphs`` in the
    same order, so positional zipping pairs each live paragraph with its stable
    anchor metadata.
    """
    anchors = dn.compute_paragraph_anchors(doc)
    paragraphs = doc.paragraphs
    for paragraph, anchor in zip(paragraphs, anchors):
        yield paragraph, anchor


def _paragraph_anchor_dict(hazard: str, anchor) -> dict:
    """Build the ``anchor`` object for a paragraph-scoped hazard."""
    return {
        "type": "paragraph",
        "hazard": hazard,
        "paragraph_key": anchor.key,
        "section": anchor.section,
        "text": _snippet(anchor.text),
    }


# --- detectors -------------------------------------------------------------


def _detect_text_boxes(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for paragraph, anchor in _iter_body_paragraph_anchors(doc):
        boxes = list(paragraph._p.iter(_q("w", "txbxContent")))
        if not boxes:
            continue
        inner = " ".join(
            "".join(t.text or "" for t in box.iter(_q("w", "t"))) for box in boxes
        )
        hazards.append(
            _Hazard(
                hazard=_TYPE_TEXT_BOX,
                severity=SEV_HIGH,
                anchor=_paragraph_anchor_dict(_TYPE_TEXT_BOX, anchor),
                current=_snippet(inner) or "(text box with no extractable text)",
                proposed="Move the text box content into normal body paragraphs.",
                rationale=(
                    "Text inside a text box (w:txbxContent) is frequently invisible "
                    "to ATS parsers, so any content placed there can be dropped entirely."
                ),
                identity=f"{_TYPE_TEXT_BOX}::{anchor.key}",
            )
        )
    return hazards


def _detect_images(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for paragraph, anchor in _iter_body_paragraph_anchors(doc):
        has_blip = next(paragraph._p.iter(_q("a", "blip")), None) is not None
        has_vml_image = next(paragraph._p.iter(_q("v", "imagedata")), None) is not None
        if not (has_blip or has_vml_image):
            continue
        hazards.append(
            _Hazard(
                hazard=_TYPE_IMAGE,
                severity=SEV_HIGH,
                anchor=_paragraph_anchor_dict(_TYPE_IMAGE, anchor),
                current=_snippet(anchor.text) or "(embedded image / drawing)",
                proposed="Replace the image with selectable text in the document body.",
                rationale=(
                    "Any text rendered inside an embedded image or drawing is not "
                    "extractable by ATS parsers and is lost during text extraction."
                ),
                identity=f"{_TYPE_IMAGE}::{anchor.key}",
            )
        )
    return hazards


def _detect_multi_column(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for section_index, section in enumerate(doc.sections):
        cols = section._sectPr.find(_q("w", "cols"))
        if cols is None:
            continue
        num = _safe_int(cols.get(_q("w", "num")))
        if num is None or num <= 1:
            continue
        hazards.append(
            _Hazard(
                hazard=_TYPE_MULTI_COLUMN,
                severity=SEV_HIGH,
                anchor={
                    "type": "section",
                    "hazard": _TYPE_MULTI_COLUMN,
                    "section_index": section_index,
                    "columns": num,
                },
                current=f"Section {section_index} uses {num} text columns.",
                proposed="Convert the section to a single-column layout.",
                rationale=(
                    "Multi-column layouts are commonly read out of order by ATS "
                    "parsers, scrambling the reading sequence of the content."
                ),
                identity=f"{_TYPE_MULTI_COLUMN}::section::{section_index}",
            )
        )
    return hazards


def _part_text(header_footer) -> str:
    element = getattr(header_footer, "_element", None)
    if element is None:
        return ""
    return "".join(t.text or "" for t in element.iter(_q("w", "t")))


def _part_name(header_footer) -> str:
    part = getattr(header_footer, "part", None)
    return str(getattr(part, "partname", "")) if part is not None else ""


def _detect_header_footer(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    seen_parts: set[str] = set()
    specs = (
        ("header", _TYPE_HEADER, "HDR"),
        ("footer", _TYPE_FOOTER, "FTR"),
    )
    for section_index, section in enumerate(doc.sections):
        for attr, hazard_type, _code in specs:
            hf = getattr(section, attr)
            # Linked parts inherit content from a previous section; only the
            # owning section should be flagged, exactly once per distinct part.
            if getattr(hf, "is_linked_to_previous", False):
                continue
            text = _part_text(hf)
            if not text.strip():
                continue
            part_name = _part_name(hf)
            dedup_key = f"{hazard_type}::{part_name or section_index}"
            if dedup_key in seen_parts:
                continue
            seen_parts.add(dedup_key)
            location = "header" if hazard_type == _TYPE_HEADER else "footer"
            hazards.append(
                _Hazard(
                    hazard=hazard_type,
                    severity=SEV_MEDIUM,
                    anchor={
                        "type": location,
                        "hazard": hazard_type,
                        "section_index": section_index,
                        "part": part_name,
                    },
                    current=_snippet(text),
                    proposed=f"Move the {location} content into the document body.",
                    rationale=(
                        f"Content placed in the document {location} is commonly "
                        "ignored by ATS parsers and may never be read."
                    ),
                    identity=f"{hazard_type}::{part_name or section_index}",
                )
            )
    return hazards


def _detect_layout_tables(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for table_index, table in enumerate(doc.tables):
        preview_cells = []
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    preview_cells.append(cell.text.strip())
            if len(preview_cells) >= 4:
                break
        preview = _snippet(" | ".join(preview_cells))
        hazards.append(
            _Hazard(
                hazard=_TYPE_LAYOUT_TABLE,
                severity=SEV_MEDIUM,
                anchor={
                    "type": "table",
                    "hazard": _TYPE_LAYOUT_TABLE,
                    "table_index": table_index,
                },
                current=preview or f"Table {table_index} (no extractable text).",
                proposed="Replace the layout table with linear body paragraphs.",
                rationale=(
                    "Tables used for visual layout are frequently read cell-by-cell "
                    "out of order by ATS parsers, fragmenting the content."
                ),
                identity=f"{_TYPE_LAYOUT_TABLE}::table::{table_index}",
            )
        )
    return hazards


def _detect_nonstandard_headings(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for paragraph, anchor in _iter_body_paragraph_anchors(doc):
        if not _is_heading_like(paragraph):
            continue
        style_name = anchor.style
        if _is_standard_heading_style(style_name):
            continue
        if not (anchor.text or "").strip():
            continue
        hazards.append(
            _Hazard(
                hazard=_TYPE_NONSTANDARD_HEADING,
                severity=SEV_MEDIUM,
                anchor=_paragraph_anchor_dict(_TYPE_NONSTANDARD_HEADING, anchor),
                current=f"{_snippet(anchor.text)} (style: {style_name or 'Normal'})",
                proposed="Apply a standard heading style (Heading 1/2/3 or Title).",
                rationale=(
                    "A heading-like paragraph that does not use a standard heading "
                    "style may not be recognised as a section boundary by ATS parsers."
                ),
                identity=f"{_TYPE_NONSTANDARD_HEADING}::{anchor.key}",
            )
        )
    return hazards


def _detect_hostile_unicode(doc) -> list[_Hazard]:
    hazards: list[_Hazard] = []
    for paragraph, anchor in _iter_body_paragraph_anchors(doc):
        chars = _hostile_chars(anchor.text)
        if not chars:
            continue
        description = _describe_chars(chars)
        hazards.append(
            _Hazard(
                hazard=_TYPE_HOSTILE_UNICODE,
                severity=SEV_MEDIUM,
                anchor=_paragraph_anchor_dict(_TYPE_HOSTILE_UNICODE, anchor),
                current=f"{_snippet(anchor.text)} [contains {description}]",
                proposed="Remove or replace the parser-hostile characters with plain equivalents.",
                rationale=(
                    "Control, format, private-use, surrogate, or unassigned characters "
                    f"({description}) commonly fail to round-trip through ATS parsers."
                ),
                identity=f"{_TYPE_HOSTILE_UNICODE}::{anchor.key}",
            )
        )
    return hazards


_DETECTORS = (
    _detect_text_boxes,
    _detect_images,
    _detect_multi_column,
    _detect_header_footer,
    _detect_layout_tables,
    _detect_nonstandard_headings,
    _detect_hostile_unicode,
)


# --- assembly --------------------------------------------------------------


def _finding_id(hazard: _Hazard) -> str:
    code = _TYPE_CODES[hazard.hazard]
    digest = hashlib.sha1(hazard.identity.encode("utf-8")).hexdigest()[:8]
    return f"ATS-{code}-{digest}"


def detect_hazards(
    doc_or_path,
    *,
    target_document: str = TARGET_CV,
    iteration: int = 0,
    source_agent: str = DEFAULT_SOURCE_AGENT,
) -> list[dict]:
    """Detect ATS structural hazards and return candidate Finding dicts.

    ``doc_or_path`` is a path (``str``/``Path``) to a ``.docx`` or an already
    opened ``python-docx`` ``Document``. Returns a deterministically ordered
    list of Findings conforming to the unified Finding schema (design "Finding"):
    each carries ``category == "ats"`` and ``status == "open"``. The list is
    empty for a structurally clean document.
    """
    if target_document not in VALID_TARGET_DOCUMENTS:
        raise ValueError(
            f"target_document must be one of {sorted(VALID_TARGET_DOCUMENTS)}; "
            f"got {target_document!r}"
        )

    if isinstance(doc_or_path, (str, Path)):
        doc = Document(str(doc_or_path))
    else:
        doc = doc_or_path

    hazards: list[_Hazard] = []
    for detector in _DETECTORS:
        hazards.extend(detector(doc))

    # Deterministic, detection-order-independent ordering: sort by (base id,
    # identity). Resolve the (astronomically unlikely) id collision by appending
    # a numeric suffix so ids stay unique within the run.
    hazards.sort(key=lambda h: (_finding_id(h), h.identity))

    findings: list[dict] = []
    used_ids: dict[str, int] = {}
    for hazard in hazards:
        base_id = _finding_id(hazard)
        count = used_ids.get(base_id, 0)
        used_ids[base_id] = count + 1
        finding_id = base_id if count == 0 else f"{base_id}-{count + 1}"
        findings.append(
            {
                "id": finding_id,
                "source_agent": source_agent,
                "iteration": iteration,
                "target_document": target_document,
                "category": CATEGORY_ATS,
                "severity": hazard.severity,
                "anchor": hazard.anchor,
                "current": hazard.current,
                "proposed": hazard.proposed,
                "rationale": hazard.rationale,
                "status": STATUS_OPEN,
            }
        )
    return findings


def build_findings_document(
    doc_or_path,
    *,
    target_document: str = TARGET_CV,
    iteration: int = 0,
    source_agent: str = DEFAULT_SOURCE_AGENT,
    source_path: Optional[str] = None,
) -> dict:
    """Wrap :func:`detect_hazards` output with schema + run metadata."""
    findings = detect_hazards(
        doc_or_path,
        target_document=target_document,
        iteration=iteration,
        source_agent=source_agent,
    )
    if source_path is None and isinstance(doc_or_path, (str, Path)):
        source_path = str(doc_or_path)
    return {
        "schema": FINDINGS_SCHEMA,
        "source_agent": source_agent,
        "iteration": iteration,
        "target_document": target_document,
        "source_document": source_path,
        "finding_count": len(findings),
        "findings": findings,
    }


def write_findings(document: dict, out_path) -> Path:
    """Write a Findings document to ``out_path`` as pretty JSON. Returns the path."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        json.dumps(document, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return out_path


# --- CLI -------------------------------------------------------------------


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        prog="ats_structural.py",
        description="Detect deterministic ATS structural hazards in a .docx and "
        "emit candidate Findings as JSON.",
    )
    parser.add_argument("input_docx", help="Path to the input .docx file.")
    parser.add_argument(
        "--out",
        default=None,
        help="Where to write the Findings JSON (default: print to stdout).",
    )
    parser.add_argument(
        "--target",
        default=TARGET_CV,
        choices=sorted(VALID_TARGET_DOCUMENTS),
        help="target_document tag applied to every emitted Finding.",
    )
    parser.add_argument(
        "--iteration",
        type=int,
        default=0,
        help="Iteration number recorded on every emitted Finding.",
    )
    parser.add_argument(
        "--source-agent",
        default=DEFAULT_SOURCE_AGENT,
        help="source_agent tag applied to every emitted Finding.",
    )
    args = parser.parse_args(argv)

    input_path = Path(args.input_docx)
    if not input_path.exists():
        sys.stderr.write(f"ERROR: input .docx not found: {input_path}\n")
        return EXIT_INPUT_NOT_FOUND

    document = build_findings_document(
        input_path,
        target_document=args.target,
        iteration=args.iteration,
        source_agent=args.source_agent,
        source_path=str(input_path),
    )

    if args.out:
        out_path = write_findings(document, args.out)
        print(
            json.dumps(
                {
                    "input_docx": str(input_path),
                    "findings_json": str(out_path),
                    "finding_count": document["finding_count"],
                }
            )
        )
    else:
        print(json.dumps(document, indent=2, ensure_ascii=False))
    return EXIT_OK


if __name__ == "__main__":
    raise SystemExit(main())
